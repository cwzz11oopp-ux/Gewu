from __future__ import annotations

from pathlib import Path
from datetime import date
import math

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
ASSET_DIR = OUT_DIR / "report_assets"
OUT_PATH = OUT_DIR / "AI科研实验系统_执行逻辑与使用指南.docx"
SCREENSHOT = Path(
    r"C:\Users\Administrator\AppData\Local\Temp"
    r"\codex-clipboard-9d863976-7ee2-46a5-8051-b3f66b82390f.png"
)

FONT_REGULAR = r"C:\Windows\Fonts\msyh.ttc"
FONT_BOLD = r"C:\Windows\Fonts\msyhbd.ttc"

NAVY = "17365D"
BLUE = "2E74B5"
CYAN = "2F75B5"
TEAL = "2A7F8E"
GREEN = "2E7D5B"
GOLD = "B88725"
RED = "A94442"
INK = "1F2937"
MUTED = "667085"
LIGHT = "F4F7FA"
LIGHT_BLUE = "E8EEF5"
LIGHT_GOLD = "FFF7E6"
LIGHT_GREEN = "EAF5EF"
WHITE = "FFFFFF"
BORDER = "CBD5E1"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def font(size: int, bold: bool = False):
    return ImageFont.truetype(FONT_BOLD if bold else FONT_REGULAR, size)


def rounded(draw, xy, fill, outline, radius=20, width=3):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def text_center(draw, box, text, fnt, fill="#17365D", line_gap=8):
    x1, y1, x2, y2 = box
    lines = text.split("\n")
    heights = []
    widths = []
    for line in lines:
        b = draw.textbbox((0, 0), line, font=fnt)
        widths.append(b[2] - b[0])
        heights.append(b[3] - b[1])
    total_h = sum(heights) + line_gap * max(0, len(lines) - 1)
    y = y1 + (y2 - y1 - total_h) / 2
    for line, w, h in zip(lines, widths, heights):
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=fnt, fill=fill)
        y += h + line_gap


def arrow(draw, start, end, color="#2E74B5", width=6):
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 18
    for delta in (2.55, -2.55):
        p = (
            end[0] + length * math.cos(angle + delta),
            end[1] + length * math.sin(angle + delta),
        )
        draw.line([end, p], fill=color, width=width)


def save_diagram(name: str, draw_fn, size=(1800, 1000)) -> Path:
    path = ASSET_DIR / name
    img = Image.new("RGB", size, "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw_fn(draw, size)
    img.save(path, quality=95)
    return path


def diagram_architecture() -> Path:
    def draw_fn(d, size):
        d.text((70, 45), "系统技术架构：界面、编排、智能体、Provider 与持久化", font=font(42, True), fill="#17365D")
        layers = [
            ("交互层", "React 19 + TypeScript + Vite\n工作台 / 设置 / 时间线 / 证据表 / 实验监控", "#EAF2F8", "#2E74B5"),
            ("API 层", "FastAPI + Pydantic\nRun / Step / Pipeline / Provider / Literature / Report API", "#E8F5F7", "#2A7F8E"),
            ("编排层", "WorkflowOrchestrator + WorkflowEngine\n步骤状态机 / 停止恢复 / 锁 / 重跑 / 校验 / 事件审计", "#EEF2FF", "#4F46A5"),
            ("智能体与 Skill 层", "Supervisor 路由 Research / Idea / Critic / Planning / Experiment / Diagnostic / Writer\nSkillRuntime = Skill 声明工具 ∩ Agent 权限 ∩ 已注册工具 ∩ 当前配置", "#FFF7E6", "#B88725"),
            ("Provider 与执行层", "Qwen OpenAI-compatible API | ArXiv + Semantic Scholar | 本地 GPU / SSH\n代码 Bundle、预检、训练进程、日志、结果 JSON、审计", "#FDEEEE", "#A94442"),
            ("数据与知识层", "JSON Repository / Artifact 谱系 / Event 审计 / Research Wiki / 本地文献库\n数据集指纹、Bundle 哈希、Attempt 目录与报告产物", "#EAF5EF", "#2E7D5B"),
        ]
        y = 135
        for i, (label, body, fill, outline) in enumerate(layers):
            h = 118 if i != 3 else 135
            rounded(d, (90, y, 1710, y + h), fill, outline, radius=22, width=3)
            rounded(d, (115, y + 21, 330, y + h - 21), "#FFFFFF", outline, radius=16, width=2)
            text_center(d, (115, y + 21, 330, y + h - 21), label, font(30, True), outline)
            text_center(d, (365, y + 10, 1680, y + h - 10), body, font(24), "#1F2937", 10)
            if i < len(layers) - 1:
                arrow(d, (900, y + h + 4), (900, y + h + 34), "#94A3B8", 5)
            y += h + 38

    return save_diagram("01_system_architecture.png", draw_fn)


def diagram_pipeline() -> Path:
    def draw_fn(d, size):
        d.text((65, 40), "九步科研 Pipeline：每一步都产生可追溯 Artifact", font=font(42, True), fill="#17365D")
        steps = [
            ("01", "问题理解", "problem", "#EAF2F8"),
            ("02", "知识整合", "evidence", "#E8F5F7"),
            ("03", "假设生成", "hypothesis", "#FFF7E6"),
            ("04", "证据推理", "reasoning / selection", "#F5EEFF"),
            ("05", "研究计划", "plan", "#EEF2FF"),
            ("06", "实验任务", "task / bundle", "#FDEEEE"),
            ("07", "运行与分析", "result / diagnosis", "#FFF1F1"),
            ("08", "反馈迭代", "revision / refined plan", "#FFF7E6"),
            ("09", "报告导出", "report", "#EAF5EF"),
        ]
        positions = []
        x0, y0 = 70, 150
        card_w, card_h, gx, gy = 500, 205, 65, 90
        for i, step in enumerate(steps):
            row, col = divmod(i, 3)
            x = x0 + col * (card_w + gx)
            y = y0 + row * (card_h + gy)
            positions.append((x, y))
            num, title, out, fill = step
            rounded(d, (x, y, x + card_w, y + card_h), fill, "#94A3B8", radius=22, width=3)
            rounded(d, (x + 22, y + 22, x + 112, y + 112), "#17365D", "#17365D", radius=45, width=2)
            text_center(d, (x + 22, y + 22, x + 112, y + 112), num, font(31, True), "#FFFFFF")
            d.text((x + 135, y + 25), title, font=font(34, True), fill="#17365D")
            d.text((x + 32, y + 132), f"输出：{out}", font=font(23), fill="#475467")
        for i in range(8):
            x, y = positions[i]
            nx, ny = positions[i + 1]
            if i % 3 != 2:
                arrow(d, (x + card_w + 8, y + card_h / 2), (nx - 8, ny + card_h / 2), "#2E74B5", 6)
            else:
                arrow(d, (x + card_w / 2, y + card_h + 8), (nx + card_w / 2, ny - 8), "#2E74B5", 6)
        d.text((90, 935), "反馈为 partial / failed 且具备可检验下一步时：08 → 新 05/06/07 → 08；达到上限或无可行动修订时停止。", font=font(24, True), fill="#A94442")

    return save_diagram("02_pipeline.png", draw_fn)


def diagram_model_routing() -> Path:
    def draw_fn(d, size):
        d.text((70, 45), "Qwen 模型路由：按任务类型分配，而不是所有步骤使用同一模型", font=font(40, True), fill="#17365D")
        rows = [
            ("Reasoning", "qwen3.7-max", "thinking=true", "假设、Idea 审阅、证据推理、语义 Reviewer、计划修订、结果分析、实验审计", "#EEF2FF", "#4F46A5"),
            ("General", "qwen-max（当前配置）", "thinking=false", "问题结构化、初始计划、报告写作等常规结构化任务", "#EAF2F8", "#2E74B5"),
            ("Code", "qwen3-coder-plus", "按模型能力决定", "实验 Bundle 生成、源码修复、诊断类代码任务；失败可回退 qwen3-coder-flash", "#FDEEEE", "#A94442"),
            ("Fast", "qwen3.6-flash", "thinking=false", "保留的低时延路由；当前主九步流程较少使用", "#EAF5EF", "#2E7D5B"),
        ]
        y = 155
        for name, model, think, tasks, fill, outline in rows:
            rounded(d, (90, y, 1710, y + 175), fill, outline, radius=22, width=3)
            rounded(d, (120, y + 28, 400, y + 147), "#FFFFFF", outline, radius=18, width=2)
            text_center(d, (120, y + 28, 400, y + 147), f"{name}\n{model}", font(27, True), outline, 8)
            d.text((450, y + 35), think, font=font(23, True), fill=outline)
            d.multiline_text((450, y + 78), tasks, font=font(24), fill="#1F2937", spacing=8)
            y += 200
        d.text((95, 920), "每次调用把 model_used、model_route、是否回退、thinking_enabled、JSON 修复状态写入事件/Artifact，便于审计。", font=font(24, True), fill="#17365D")

    return save_diagram("03_model_routing.png", draw_fn)


def diagram_evidence() -> Path:
    def draw_fn(d, size):
        d.text((70, 40), "证据推理：先做确定性证据审计，再允许模型进行语义判断", font=font(40, True), fill="#17365D")
        boxes = [
            ("1  证据注册表", "对 DOI / arXiv / URL / 标题去重，生成稳定 evidence_id；只保留验证状态与来源。", "#EAF2F8", "#2E74B5"),
            ("2  候选级审计", "逐个候选拆分主张；建立 candidate_audit，预检查可匹配证据与硬门槛。", "#E8F5F7", "#2A7F8E"),
            ("3  Idea Reviewer", "对每个候选给出 evidence ledger、closest prior work、gates、scores、MDE、decision。JSON 形状最多纠正 3 次。", "#FFF7E6", "#B88725"),
            ("4  Critic 逐候选推理", "必须返回 claim_evidence_map；每条关系标注 support/contradict/context 与 DIRECT/INDIRECT/ANALOGY。", "#F5EEFF", "#6F42A5"),
            ("5  系统硬门槛", "发明的 evidence_id、只有类比证据、缺少直接/间接支持、candidate gate=FAIL 时，不允许 GO。", "#FDEEEE", "#A94442"),
            ("6  一次定向补证", "若没有可行候选，提取最多 3 个缺口查询；只执行一轮检索，合并新增已验证证据后重新审阅。", "#EAF5EF", "#2E7D5B"),
            ("7  选择与留痕", "仅从 verified/revised 候选中按审阅结果选优；保存全部候选评估、证据注册表、策略与选择理由。", "#EEF2FF", "#4F46A5"),
        ]
        y = 125
        for i, (title, body, fill, outline) in enumerate(boxes):
            rounded(d, (140, y, 1660, y + 105), fill, outline, radius=18, width=3)
            d.text((180, y + 20), title, font=font(27, True), fill=outline)
            d.text((580, y + 20), body, font=font(22), fill="#1F2937")
            if i < len(boxes) - 1:
                arrow(d, (900, y + 107), (900, y + 132), "#94A3B8", 5)
            y += 128

    return save_diagram("04_evidence_reasoning.png", draw_fn)


def diagram_experiment_loop() -> Path:
    def draw_fn(d, size):
        d.text((70, 42), "实验执行与有界诊断修复：快门槛优先，正式训练最后", font=font(40, True), fill="#17365D")
        nodes = [
            (100, 170, 440, 300, "Plan → Task → Bundle\n固定 dataset / parameters / seeds / metrics", "#EEF2FF", "#4F46A5"),
            (570, 170, 910, 300, "静态与协议预检\n语法、依赖、API、参数、数据契约、CUDA", "#EAF2F8", "#2E74B5"),
            (1040, 170, 1380, 300, "Smoke Test\nforward / loss / backward / serialize", "#E8F5F7", "#2A7F8E"),
            (1040, 450, 1380, 580, "正式执行\n本地 GPU 或 SSH；append-only attempt", "#FFF7E6", "#B88725"),
            (570, 450, 910, 580, "Result + Analysis + Audit\n只读结果文件，不从 stdout 猜指标", "#EAF5EF", "#2E7D5B"),
            (100, 450, 440, 580, "诊断分类\n数据 / 依赖 / GPU / 代码 / 超时 / 审计", "#FDEEEE", "#A94442"),
            (100, 735, 440, 865, "最小修复\n隔离坏缓存、重试阶段、修源码、必要时重建 Bundle", "#FFF1F1", "#A94442"),
            (570, 735, 910, 865, "边界\n运行修复最多 2 轮；代码候选每轮最多 3 个", "#FFF7E6", "#B88725"),
            (1040, 735, 1380, 865, "失败落盘\n仍生成 failed result + attempts + diagnosis", "#F4F7FA", "#667085"),
        ]
        for x1, y1, x2, y2, label, fill, outline in nodes:
            rounded(d, (x1, y1, x2, y2), fill, outline, radius=20, width=3)
            text_center(d, (x1 + 15, y1 + 10, x2 - 15, y2 - 10), label, font(22, True), "#1F2937", 8)
        for a, b in [((440,235),(570,235)),((910,235),(1040,235)),((1210,300),(1210,450)),((1040,515),(910,515)),((570,515),(440,515)),((270,580),(270,735)),((440,800),(570,800))]:
            arrow(d, a, b, "#2E74B5", 6)
        arrow(d, (910, 800), (1210, 580), "#B88725", 6)
        arrow(d, (570, 765), (390, 300), "#A94442", 6)
        d.text((1450, 185), "预检通过", font=font(23, True), fill="#2E7D5B")
        d.text((1450, 470), "执行成功", font=font(23, True), fill="#2E7D5B")
        d.text((1450, 755), "不可修复", font=font(23, True), fill="#A94442")

    return save_diagram("05_experiment_loop.png", draw_fn)


def diagram_lineage() -> Path:
    def draw_fn(d, size):
        d.text((70, 45), "Artifact 谱系与“重新运行本次实验”的正确语义", font=font(41, True), fill="#17365D")
        nodes = [
            (80, 190, "Plan v1", "#EAF2F8", "#2E74B5"),
            (340, 190, "Task v1", "#EEF2FF", "#4F46A5"),
            (600, 190, "Bundle v1", "#FDEEEE", "#A94442"),
            (860, 190, "Result v1", "#FFF7E6", "#B88725"),
            (1120, 190, "Revision 1", "#F5EEFF", "#6F42A5"),
            (1380, 190, "Plan v2", "#EAF5EF", "#2E7D5B"),
            (1380, 500, "Task v2", "#EEF2FF", "#4F46A5"),
            (1120, 500, "Bundle v2", "#FDEEEE", "#A94442"),
            (860, 500, "Result v2\nattempt 1", "#FFF7E6", "#B88725"),
            (600, 500, "Result v2\nattempt 2", "#EAF5EF", "#2E7D5B"),
        ]
        boxes = {}
        for x, y, label, fill, outline in nodes:
            boxes[label] = (x, y, x + 220, y + 125)
            rounded(d, boxes[label], fill, outline, radius=18, width=3)
            text_center(d, boxes[label], label, font(25, True), "#1F2937", 6)
        seq1 = [nodes[i] for i in range(6)]
        for a, b in zip(seq1, seq1[1:]):
            arrow(d, (a[0] + 220, a[1] + 62), (b[0], b[1] + 62), "#64748B", 5)
        arrow(d, (1490, 315), (1490, 500), "#64748B", 5)
        for a, b in [(nodes[6], nodes[7]), (nodes[7], nodes[8]), (nodes[8], nodes[9])]:
            arrow(d, (a[0], a[1] + 62), (b[0] + 220, b[1] + 62), "#64748B", 5)
        rounded(d, (210, 760, 1590, 920), "#F4F7FA", "#2E74B5", radius=22, width=3)
        text_center(
            d,
            (240, 775, 1560, 905),
            "重跑规则：保留 Revision 1 → Plan v2 → Task v2 → Bundle v2 全部祖先谱系；\n"
            "保留旧 Result v2 作为历史，只为同一 Bundle 新增 attempt；仅使旧报告失效。\n"
            "事件中记录 task_artifact_id、bundle_artifact_id、previous_result_artifact_id。",
            font(26, True),
            "#17365D",
            12,
        )

    return save_diagram("06_lineage_retry.png", draw_fn)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=TABLE_INDENT_DXA):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
    set_table_borders(table)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=11, color=INK, bold=False, italic=False, font_name="Calibri"):
    run.font.name = font_name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def set_paragraph_spacing(p, before=0, after=6, line=1.25):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.22

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 9),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 11.5, NAVY, 10, 5),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = rgb(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for sec in doc.sections:
        header = sec.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p, after=0, line=1)
        r = p.add_run("AI 科研实验系统 · 技术执行逻辑与操作指南")
        set_run_font(r, size=8.5, color=MUTED, bold=True)
        footer = sec.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = fp.add_run("内部技术说明  |  ")
        set_run_font(r, size=8, color=MUTED)
        add_page_number(fp)


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=0, after=8, line=1)
    r = p.add_run(text)
    set_run_font(r, size=28, color=NAVY, bold=True)
    if subtitle:
        sp = doc.add_paragraph()
        set_paragraph_spacing(sp, after=18, line=1.15)
        r = sp.add_run(subtitle)
        set_run_font(r, size=13.5, color=MUTED)


def add_kicker(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6, line=1)
    r = p.add_run(text.upper())
    set_run_font(r, size=9, color=GOLD, bold=True)


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    set_paragraph_spacing(p)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True, color=NAVY)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.38 + level * 0.24)
    p.paragraph_format.first_line_indent = Inches(-0.19)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(text)
    set_run_font(r, size=10.2)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(text)
    set_run_font(r, size=10.3)
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE, color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_cell_shading(table.cell(0, 0), fill)
    p = table.cell(0, 0).paragraphs[0]
    set_paragraph_spacing(p, before=2, after=2, line=1.2)
    r = p.add_run(f"{label}  ")
    set_run_font(r, size=10.3, color=color, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.3, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = Pt(1)
    set_run_font(spacer.add_run("\u200b"), size=1, color="FFFFFF")


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, after=0, line=1.12)
        r = p.add_run(header)
        set_run_font(r, size=font_size, color=NAVY, bold=True)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        if row_index % 2:
            for cell in cells:
                set_cell_shading(cell, "F9FBFD")
        for i, value in enumerate(values):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(p, after=0, line=1.12)
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=INK, bold=(i == 0))
    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = Pt(1)
    set_run_font(spacer.add_run("\u200b"), size=1, color="FFFFFF")
    return table


def add_figure(doc, path: Path, caption: str, width=6.55):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=3, after=4, line=1)
    picture = p.add_run().add_picture(str(path), width=Inches(width))
    picture._inline.docPr.set("title", caption)
    picture._inline.docPr.set("descr", caption)
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(cp, before=0, after=8, line=1.12)
    r = cp.add_run(caption)
    set_run_font(r, size=8.8, color=MUTED, italic=True)


def section_title(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = True
    set_paragraph_spacing(p, before=0, after=6, line=1)
    r = p.add_run("SYSTEM EXECUTION DETAIL")
    set_run_font(r, size=9, color=GOLD, bold=True)
    h = doc.add_paragraph(style="Heading 1")
    r = h.add_run(title)
    set_run_font(r, size=18, color=NAVY, bold=True)
    if subtitle:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=12)
        r = p.add_run(subtitle)
        set_run_font(r, size=11, color=MUTED, italic=True)


def add_step_chapter(doc, step):
    section_title(doc, f"{step['index']}  {step['title']}", step["subtitle"])
    add_callout(
        doc,
        "步骤定位",
        f"Agent：{step['agent']}；Skill：{'、'.join(step['skills'])}；模型路由：{step['model']}；核心输出：{step['output']}。",
    )
    doc.add_paragraph("输入与前置条件", style="Heading 2")
    for item in step["inputs"]:
        add_bullet(doc, item)
    doc.add_paragraph("执行逻辑", style="Heading 2")
    for item in step["logic"]:
        add_number(doc, item)
    doc.add_paragraph("输出与谱系", style="Heading 2")
    for item in step["outputs"]:
        add_bullet(doc, item)
    doc.add_paragraph("校验、失败与恢复", style="Heading 2")
    for item in step["validation"]:
        add_bullet(doc, item)
    if step.get("note"):
        add_callout(doc, "实现说明", step["note"], fill=LIGHT_GOLD, color=GOLD)


def build_report():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = {
        "architecture": diagram_architecture(),
        "pipeline": diagram_pipeline(),
        "models": diagram_model_routing(),
        "evidence": diagram_evidence(),
        "experiment": diagram_experiment_loop(),
        "lineage": diagram_lineage(),
    }

    doc = Document()
    configure_document(doc)

    # Cover
    add_kicker(doc, "TECHNICAL REPORT · 2026")
    for _ in range(3):
        doc.add_paragraph()
    add_title(
        doc,
        "AI 科研实验系统\n执行逻辑、技术架构与使用指南",
        "逐步解释数据集绑定、文献证据、假设生成、证据推理、实验执行、诊断修复、反馈迭代与报告导出",
    )
    add_callout(
        doc,
        "报告范围",
        "依据当前 D:\\竞赛 工作区的实际代码、Skill 路由和运行配置整理。模型密钥、SSH 密钥及其他敏感配置均不写入本报告。",
        fill=LIGHT_GREEN,
        color=GREEN,
    )
    add_table(
        doc,
        ["项目", "当前值"],
        [
            ("文档类型", "技术白皮书 + 操作手册"),
            ("系统流程", "9 个主步骤 + 实验诊断支路 + 反馈迭代闭环"),
            ("当前 LLM", "Qwen 多模型路由"),
            ("当前文献源", "Research Wiki / 本地文献 / ArXiv / Semantic Scholar"),
            ("当前实验环境", "本地 GPU，CUDA_VISIBLE_DEVICES=0"),
            ("当前数据集", "本地 IPIX，启动前检查并绑定内容指纹"),
            ("生成日期", str(date(2026, 7, 24))),
        ],
        [2100, 7260],
        9.5,
    )
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=12, after=0)
    r = p.add_run("适用读者：系统使用者、竞赛成员、实验复核人员、后续开发与维护人员。")
    set_run_font(r, size=10.5, color=MUTED, italic=True)

    # Overview
    section_title(doc, "1  报告阅读指南与核心结论")
    add_body(doc, "本系统不是单次 Prompt 调用，而是一个带状态、Artifact 谱系、Skill 权限、模型路由、验证循环和实验运行时的科研工作流。每一步只读取被允许的上游产物，产生新 Artifact，并把模型、Skill、工具、错误、修订和来源关系写入事件记录。")
    add_callout(
        doc,
        "一句话理解",
        "系统把“问题 → 文献证据 → 候选假设 → 证据审查 → 实验计划 → 可执行代码 → 真实结果 → 反馈修订 → 报告”拆成九个可暂停、可审计、可重跑的状态转换。",
    )
    doc.add_paragraph("建议阅读顺序", style="Heading 2")
    for item in [
        "首次使用者：先看第 2、3、16、17 章，理解总体流程和操作方法。",
        "关注证据推理：重点看第 8 章和“证据推理深潜”。",
        "关注实验失败与迭代：重点看第 11、12、13、14 章。",
        "开发维护者：重点看技术栈、模型路由、SkillRuntime、Artifact 数据模型与 API 附录。",
    ]:
        add_bullet(doc, item)
    doc.add_paragraph("当前关键设计结论", style="Heading 2")
    for item in [
        "证据推理由“确定性证据审计 + LLM 语义评审”共同完成，模型不能自行发明 evidence_id，也不能把类比证据写成直接证据。",
        "模型按任务路由：深度推理、常规生成、代码生成和快速任务使用不同 Qwen 模型与超时策略。",
        "本地数据集在问题理解前被检查、生成内容指纹并锁定；后续计划与 Bundle 不允许静默更换数据集。",
        "实验失败不是流程丢失：系统保存 failed result、attempts 和 diagnosis，然后决定代码修复还是科学计划修订。",
        "重新运行本次实验现在绑定当前迭代的 task/bundle，只新增 attempt，不再回退到初始实验。",
    ]:
        add_bullet(doc, item)

    section_title(doc, "2  系统总体架构")
    add_figure(doc, diagrams["architecture"], "图 1  系统分层架构与主要技术组件")
    add_body(doc, "前端只负责交互与可视化；真正的科研状态保存在后端 Repository 中。WorkflowOrchestrator 负责自动推进，WorkflowEngine 负责单步原子执行。Supervisor 不直接做领域工作，而是静态路由到对应 Agent 和 Skill，并在产物入库前执行结构与语义验证。")
    architecture_detail = doc.add_paragraph(style="Heading 2")
    architecture_detail.paragraph_format.page_break_before = True
    architecture_detail.add_run("分层职责详表")
    add_table(
        doc,
        ["层级", "主要技术", "职责"],
        [
            ("交互层", "React 19、TypeScript、Vite、Lucide", "创建运行、设置 Provider、展示时间线、Artifact、证据和实验进度"),
            ("API 层", "FastAPI、Pydantic、HTTP JSON", "Run/Step/Pipeline/Literature/Settings/Report 等接口"),
            ("编排层", "线程式 Orchestrator、WorkflowEngine、RLock", "状态机、原子步骤、停止、恢复、重跑、验证和追踪"),
            ("Agent 层", "7 类领域 Agent + Supervisor + Reviewer", "结构化问题、文献、Idea、批评、计划、实验、诊断、报告"),
            ("Skill 层", "本地 SKILL.md、静态路由、工具交集授权", "提供操作协议、边界、输出合同和可调用工具"),
            ("Provider 层", "Qwen、ArXiv、Semantic Scholar、本地 GPU/SSH", "模型推理、外部检索、实验执行"),
            ("数据层", "JSON Store、Artifact、Event、Wiki、文献库", "持久化状态、谱系、证据来源、可恢复历史"),
        ],
        [1250, 2600, 5510],
        8.7,
    )

    section_title(doc, "3  九步主流程总览")
    add_figure(doc, diagrams["pipeline"], "图 2  九步科研 Pipeline 与主要 Artifact")
    add_table(
        doc,
        ["步骤", "Agent", "固定 Skill", "主要 Artifact"],
        [
            ("1 问题理解", "Research", "problem-framing", "dataset_profile、problem"),
            ("2 知识整合", "Research", "research-lit、research-wiki", "evidence、Wiki changes"),
            ("3 假设生成", "Idea", "idea-creator", "hypothesis candidates"),
            ("4 证据推理", "Critic", "idea-selection、novelty-check、research-review", "idea_review、selection、reasoning"),
            ("5 研究计划", "Planning", "research-refine、hypothesis-experiment-gate、experiment-plan", "plan"),
            ("6 实验任务", "Experiment", "experiment-implementation", "experiment_task、experiment_bundle"),
            ("7 运行分析", "Experiment/Diagnostic", "run-experiment、analyze-results、experiment-audit；失败时 experiment-diagnosis", "result、diagnosis、repaired bundle"),
            ("8 反馈迭代", "Critic", "experiment-iteration、result-to-claim；条件触发 ablation-planner 等", "revision、refined plan"),
            ("9 报告导出", "Writer", "competition-report", "report"),
        ],
        [1180, 1350, 3600, 3230],
        8.2,
    )

    section_title(doc, "4  模型路由与当前模型配置")
    add_figure(doc, diagrams["models"], "图 3  Qwen 多模型路由策略")
    add_callout(
        doc,
        "当前运行证据",
        "已检查的 IPIX 运行事件中，深度推理路由实际记录 qwen3.7-max；常规路由实际记录 qwen-max。代码路由的默认主模型为 qwen3-coder-plus，回退模型为 qwen3-coder-flash。",
        fill=LIGHT_GREEN,
        color=GREEN,
    )
    add_table(
        doc,
        ["路由", "主模型", "回退", "Thinking", "典型任务"],
        [
            ("reasoning", "qwen3.7-max", "general model", "开启", "假设、Idea 评审、证据推理、Reviewer、计划修订、结果分析、审计"),
            ("general", "qwen-max（当前）", "reasoning model", "关闭", "问题结构化、初始计划、报告等"),
            ("code", "qwen3-coder-plus", "qwen3-coder-flash", "由模型兼容性决定", "Bundle 生成、代码修复、诊断"),
            ("fast", "qwen3.6-flash", "无", "关闭", "低时延预留任务"),
        ],
        [1100, 1900, 1900, 1500, 2960],
        8.4,
    )
    add_body(doc, "QwenProvider 通过 OpenAI-compatible Chat Completions API 调用模型；TaskPolicy 决定模型、thinking、超时与回退。系统对 400、404、408、409、429、500、502、503、504 默认额外重试 1 次，并校验结构化输出的必需字段；多层 JSON 仅在能唯一识别时解包。Event 保存实际模型、路由、回退、thinking、JSON 修复与 shape 规范化元数据。")

    section_title(doc, "5  SkillRuntime、Agent 与工具权限")
    add_body(doc, "Skill 不是一个额外模型，而是一份本地执行协议。每个 SKILL.md 声明名称、用途、允许工具、输入/输出合同、禁止事项和验收条件。Supervisor 使用静态表把步骤映射到 Agent 与 Skill；SkillRuntime 再将完整 Skill 指令注入该步骤。")
    add_callout(
        doc,
        "最终工具权限公式",
        "授权工具 = Skill 声明工具 ∩ Agent 工具白名单 ∩ ToolRegistry 已注册工具 ∩ 当前环境已配置工具。任何一项不满足都会得到 SKILL_TOOL_UNAUTHORIZED。",
    )
    add_table(
        doc,
        ["Agent", "主要职责", "可用工具类型"],
        [
            ("Research", "问题结构化、Wiki 和文献检索", "query_wiki、search_local_literature、literature_search"),
            ("Idea", "生成候选假设", "read_run、read_artifact、read_wiki_query_pack"),
            ("Critic", "证据审查、结果到主张判断", "literature_search、audit_evidence、audit_result"),
            ("Planning", "实验计划与反馈计划修订", "read_run、read_artifact"),
            ("Experiment", "Bundle、执行、结果读取与审计", "build_experiment_bundle、local_process_run、ssh_run、read_experiment_result"),
            ("Diagnostic", "故障分类与受限修复", "repair_dataset_cache、retry_experiment、build_experiment_bundle"),
            ("Writer", "从合格 Artifact 生成报告", "render_report"),
            ("Supervisor", "路由、校验、请求修订、事件与 Wiki 提交", "dispatch_agent、validate_artifact、request_revision 等"),
        ],
        [1250, 3700, 4410],
        8.7,
    )
    add_body(doc, "Skill 指令有 32,000 字符预算，并为每个 Skill 和合并指令计算 SHA-256。事件里记录 skill_invocations、instruction_sha256、authorized_tools、denied_tools 和 omitted_sections，因此可以复核“当时模型实际收到了哪套规则”。")

    section_title(doc, "6  Artifact、Event 与状态机")
    add_body(doc, "RunRecord 是一次科研运行的聚合根，包含状态、当前步骤、自动模式、停止标志、反馈轮次、强制新 attempt 标志、StepRecord、Artifact 和 Event。Artifact 是不可替代的科研产物，Event 是发生过的动作与审计日志。")
    add_table(
        doc,
        ["对象", "关键字段", "作用"],
        [
            ("RunRecord", "id、status、current_step、automatic、stop_requested", "表示整次工作流的当前状态"),
            ("StepRecord", "status、input/output ids、started/completed、error", "表示单个步骤的运行状态"),
            ("Artifact", "id、type、content、source_step、agent、parent_artifact_id、locked", "保存结构化产物与父子谱系"),
            ("EventRecord", "timestamp、step、actor、message、data、output_summary、tool_calls", "保存可审计事件与模型/工具元数据"),
            ("Experiment attempt", "attempt、start/end、status、error_code、log_path、recovered", "追加式保存每次实际执行"),
        ],
        [1500, 3900, 3960],
        8.8,
    )
    add_body(doc, "步骤运行前先把状态置为 running；成功后置 completed；异常时保存 code/message 并置 failed；用户停止会抛出 LLMRequestCancelled，步骤置 interrupted，Run 置 paused。每个 run 使用 RLock，避免同一运行同时执行两个互相冲突的步骤。")
    add_body(doc, "锁定 Artifact 后，从同一步再次运行会跳过该步骤；但反馈步骤会额外确认锁定 revision 是否属于最新 result，避免旧反馈阻止新结果复核。")

    steps = [
        {
            "index": "7",
            "title": "步骤 1：问题理解（problem_understanding）",
            "subtitle": "把自然语言研究要求转换为可检验的问题合同，并在最前端锁定本地数据集。",
            "agent": "Research Agent",
            "skills": ["problem-framing"],
            "model": "general（当前 qwen-max）",
            "output": "dataset_profile、problem",
            "inputs": [
                "用户输入的研究问题、领域与约束；Run 处于可执行状态。",
                "实验 Provider 的数据集设置。若 source=local，则目录必须存在且可检查。",
            ],
            "logic": [
                "若选择本地数据集，Dataset Inspector 先解析目录，统计文件数、类型、大小与样例 Schema。",
                "根据目录内容生成 contract_id 和 content_fingerprint；将 dataset_profile 保存并锁定。",
                "把“本地数据集为权威事实、禁止替换为公开/合成/其他数据集”的约束追加到 Skill 指令。",
                "调用 ResearchAgent.structure_problem，让 general 模型输出 problem_statement、constraints、knowledge_gaps、literature_queries。",
                "把数据集 profile 嵌入 problem，以便所有下游步骤继承同一个数据契约。",
            ],
            "outputs": [
                "dataset_profile：本地路径、文件数、类型、Schema、指纹、检查状态。",
                "problem：可证伪问题、约束、知识缺口和精确检索式。",
                "Event：数据检查事件和“Structured problem input”模型调用记录。",
            ],
            "validation": [
                "本地目录为空、缺失或路径未配置时立即失败，不允许用别的数据集兜底。",
                "Supervisor 验证 problem_statement、constraints、knowledge_gaps、literature_queries 必须存在。",
                "常规验证最多请求模型修订 2 次；仍不合格则步骤失败并保存原因。",
            ],
            "note": "",
        },
        {
            "index": "8",
            "title": "步骤 2：知识整合（knowledge_integration）",
            "subtitle": "先查本地知识，再查外部来源；去重、验证并保存可追溯证据卡。",
            "agent": "Research Agent",
            "skills": ["research-lit", "research-wiki"],
            "model": "主要为检索与确定性整合；必要时 general",
            "output": "evidence、WikiChangeSet",
            "inputs": [
                "步骤 1 的 problem，尤其是完整 literature_queries。",
                "Research Wiki、本地文献库、ArXiv 与 Semantic Scholar Provider。",
            ],
            "logic": [
                "对每个标准化查询先访问 Research Wiki；空 Wiki、无匹配或降级不是致命错误。",
                "随后搜索本地文献库，再访问 ArXiv 与 Semantic Scholar。",
                "按 DOI → arXiv ID → 规范化标题的顺序去重。",
                "保留作者、年份、URL、identifier、来源类型、验证 Provider、支持的主张与 exportable 状态。",
                "将本地-only 记录与外部已验证引用分开；上传文件本身不能自动成为可导出引用。",
                "生成 WikiChangeSet，只有 Supervisor 可以提交论文、Gap 与边的变更。",
            ],
            "outputs": [
                "evidence.references：已验证且可追溯的 EvidenceCard。",
                "local_only、warnings、sources.calls：不满足外部引用条件的材料和每次检索调用。",
                "Research Wiki 更新及“Collected N verified references”事件。",
            ],
            "validation": [
                "Provider 单点失败只形成 warning，其他来源的证据仍保留。",
                "不得推断缺失作者、年份或书目信息。",
                "references 必须是列表；内容会在后续证据审计再次过滤。",
            ],
        },
        {
            "index": "9",
            "title": "步骤 3：假设生成（hypothesis_generation）",
            "subtitle": "从问题、证据和 Wiki 查询包生成少量真正不同、可以证伪的候选。",
            "agent": "Idea Agent",
            "skills": ["idea-creator"],
            "model": "reasoning（qwen3.7-max，thinking 开启）",
            "output": "hypothesis candidates",
            "inputs": [
                "problem、EvidenceCard 列表和只读 Wiki query pack。",
                "本地数据集 profile（若存在），作为强制数据绑定。",
            ],
            "logic": [
                "向模型提供问题合同、证据摘要和 Skill 指令。",
                "要求候选包含 claim、机制链、与最近工作的差异、可证伪预测、失败条件、数据/算力/指标和不确定性。",
                "若为本地数据集，指令明确禁止候选提出其他数据集或合成替代。",
                "标准化候选结构并检查非空；当前策略偏好 3 个技术上不同的候选，而不是同义改写。",
                "不在本步骤选赢家，也不允许在生成过程中偷偷执行新搜索。",
            ],
            "outputs": [
                "hypothesis Artifact：candidates 数组及模型路由元数据。",
                "Event：候选生成、证据数量、模型与 Skill 哈希。",
            ],
            "validation": [
                "空候选、候选数不足或字段不完整会进入 Supervisor 修订循环。",
                "若最终没有可用候选，保存原始输出摘要并抛出 HYPOTHESIS_CANDIDATES_EMPTY。",
                "Mock fallback 仅用于开发，不能作为竞赛科学推理。",
            ],
        },
        {
            "index": "10",
            "title": "步骤 4：证据推理（evidence_reasoning）",
            "subtitle": "CCFA 风格的 Idea 审阅已增强为“确定性证据审计 + 逐候选模型批评 + 硬门槛”。",
            "agent": "Critic Agent",
            "skills": ["idea-selection", "novelty-check", "research-review"],
            "model": "reasoning（qwen3.7-max，thinking 开启）",
            "output": "idea_review、hypothesis_selection、reasoning",
            "inputs": [
                "problem、所有 verified evidence、标准化候选假设。",
                "EvidenceAudit 生成的 registry、candidate_audits 和 policy。",
            ],
            "logic": [
                "确定性构建 evidence registry，为每条证据生成稳定 evidence_id。",
                "对每个候选生成候选级审计，准备可引用 ID、匹配情况和硬门槛。",
                "IdeaSelectionAgent 一次审阅所有候选；若 JSON 形状错误，最多单独纠正 3 次格式。",
                "CriticAgent 再逐候选输出 claim_evidence_map，明确 stance 与 DIRECT/INDIRECT/ANALOGY。",
                "系统校验 evidence_id 必须存在；类比证据只能说明机制可行，不能当目标任务直接性能证据。",
                "candidate gate=FAIL、证据 ID 伪造或缺少直接/间接支持时，模型高分也不能覆盖硬门槛。",
                "若没有 verified/revised 候选，系统提取最多 3 个证据缺口，只执行一轮定向检索，再从头重评。",
                "从合格候选中选择最高评价者，保留原始/修订假设、所有候选审阅与选择理由。",
            ],
            "outputs": [
                "idea_review：每个候选的 evidence ledger、closest prior work、gates、scores、MDE、risks、decision。",
                "hypothesis_selection：selected index、选择模式与审阅结论。",
                "reasoning：active_hypothesis、candidate_assessments、evidence_registry、evidence_policy、targeted_retrieval。",
            ],
            "validation": [
                "没有合格候选且定向补证后仍失败时，输出 EVIDENCE_INSUFFICIENT，而不是强行选择。",
                "每个候选必须被审阅和推理，不能只解释最终赢家。",
                "只允许一轮定向补证，防止无界搜索和不可控成本。",
            ],
        },
        {
            "index": "11",
            "title": "步骤 5：研究计划（research_plan）",
            "subtitle": "把选中假设转换为最小、可复现、能够确认或证伪主张的实验合同。",
            "agent": "Planning Agent",
            "skills": ["research-refine", "hypothesis-experiment-gate", "experiment-plan"],
            "model": "初始 general；反馈修订 reasoning",
            "output": "plan",
            "inputs": [
                "active_hypothesis、problem、数据集选项、资源约束和证据结论。",
                "本地 dataset_profile 或在线/缓存数据集目录。",
            ],
            "logic": [
                "Hypothesis Experiment Gate 先区分 hypothesis_falsified、experiment_invalid 和 hypothesis_underspecified。",
                "把主张拆成 alignment contract：对象、任务、数据、基线、变体、指标与阈值必须一致。",
                "设计最小实验矩阵，明确唯一变化因素、固定控制、强基线、seeds、统计摘要和停止条件。",
                "为每个实验写明为什么具有诊断性、证据依据、成功/失败阈值及正负解释。",
                "本地模式强制把 plan.dataset 绑定到 contract_id 与 content_fingerprint。",
                "附加 dataset card；若在线数据集不可下载/缓存，则拒绝计划。",
            ],
            "outputs": [
                "plan：objective、hypotheses、method、dataset、comparisons、evaluations、procedure、parameters、seeds、criteria、resources、risks。",
                "claim-to-metric traceability：每个主张如何被具体比较和指标检验。",
            ],
            "validation": [
                "未命名干预、缺少基线、同时改变多个因果因素、无决策阈值、超预算或数据集冲突均拒绝。",
                "Supervisor 结构验证后，Reviewer 还会做 falsifiability、procedure completeness、resource feasibility 语义复核。",
                "常规候选最多修订 2 次。",
            ],
        },
        {
            "index": "12",
            "title": "步骤 6：实验任务与 Bundle（experiment_task）",
            "subtitle": "用代码模型生成 Provider-neutral 的完整实验包，并由后端补齐不可变协议字段。",
            "agent": "Experiment Agent",
            "skills": ["experiment-implementation"],
            "model": "code（qwen3-coder-plus；回退 qwen3-coder-flash）",
            "output": "experiment_task、experiment_bundle",
            "inputs": [
                "最新 plan 及 dataset card。",
                "执行 Provider 的 Python 命令与运行能力。",
            ],
            "logic": [
                "ExperimentAgent 先生成 task：experiment_id、result_id、manifest 与科学合同摘要。",
                "代码模型只返回一个 train.py 的 content_lines 和 requirements；后端确定性添加入口、IDs、output 与参数。",
                "验证路径必须安全且相对；每个 content_lines 元素必须恰好是一行物理 Python。",
                "源码必须接受 --run-id、--experiment-id、--result-id、--output，并写出包含全部预期 metrics 的有限 JSON。",
                "baseline 与 variant 只能改变被检验因素；禁止用不同变量名实现同一个模型。",
                "若有 epoch 循环，必须输出 JSONL progress 事件，供实时监控读取。",
            ],
            "outputs": [
                "experiment_task：实验标识、manifest 摘要及 result_id。",
                "experiment_bundle：完整源码、requirements、参数、seeds、GPU 要求和 expected metrics。",
                "谱系：task.parent = plan；bundle.parent = task。",
            ],
            "validation": [
                "compile、协议字段、依赖声明、指标来源、数据集契约和实验差异都会被静态检查。",
                "候选代码不合格时把具体错误反馈给模型重新生成，最多 2 次常规修订。",
                "本步骤不安装依赖、不下载数据、不执行训练。",
            ],
        },
        {
            "index": "13",
            "title": "步骤 7：实验运行、分析与审计（experiment_run_analysis）",
            "subtitle": "先恢复或创建 attempt，再执行、读取结果、分析并独立审计；失败进入有界修复。",
            "agent": "Experiment Agent + Experiment Diagnostic Agent",
            "skills": ["run-experiment", "analyze-results", "experiment-audit", "失败时 experiment-diagnosis", "可选 monitor-experiment"],
            "model": "执行为本地进程；分析/审计 reasoning；诊断/修码 code",
            "output": "experiment_result、experiment_diagnosis、可选 repaired bundle",
            "inputs": [
                "最新 plan、task、bundle，以及同实验旧 result/attempts（若存在）。",
                "本地 GPU 或 SSH 配置、数据集目录和 Python 环境。",
            ],
            "logic": [
                "若存在完全匹配 task、bundle、hash 与 schema 的已完成结果，且未强制新 attempt，则恢复而不重复训练。",
                "否则创建新的追加式 attempt；先做语法、依赖、API、数据、参数和 CUDA 预检，再做 smoke test。",
                "通过后启动本地 GPU/SSH 进程；日志中的 JSONL progress 只用于监控，不作为最终指标来源。",
                "只从声明的 result JSON 读取 metrics，并验证 run/experiment/result IDs、有限数值和预期字段。",
                "Analyze Results 按父 plan 的阈值做比较；Audit 再核对源码哈希、命令、环境、shape、泄漏、baseline/variant 一致性和种子。",
                "任一阶段抛错时生成 diagnosis，分类并决定是否允许确定性修复。",
                "修复后重新验证；最多 2 个运行时修复循环。最终失败也保存 failed result、attempts、error 和 diagnosis。",
            ],
            "outputs": [
                "experiment_result：原始 metrics、analysis、audit、environment、attempts、is_real_experiment。",
                "experiment_diagnosis：category、error_code、root_cause、evidence、repair_action、user_message、next_action。",
                "若修复代码，保存新的 repaired bundle 并接入原任务谱系。",
            ],
            "validation": [
                "进程成功不等于科学结果合格；必须同时通过结果 schema 和 audit。",
                "GPU 必需时 CUDA 探针失败会使 is_real_experiment=false。",
                "不能自动安装依赖、改驱动、改凭据、改科学计划或弱化校验。",
                "失败的新 attempt 也会消费强制重跑标志，避免污染后续迭代。",
            ],
        },
        {
            "index": "14",
            "title": "步骤 8：结果反馈与迭代（feedback_revision）",
            "subtitle": "从审计结果判断主张，不把执行成功误当作科学支持；只安排能改变结论的最小下一步。",
            "agent": "Critic Agent",
            "skills": ["experiment-iteration", "result-to-claim", "条件触发 research-refine / experiment-plan / ablation-planner"],
            "model": "reasoning（qwen3.7-max）",
            "output": "revision；必要时 refined plan",
            "inputs": [
                "最新 experiment_result 及其精确 plan/task/bundle/hypothesis 谱系。",
                "analysis、audit、failed attempts、指标、阈值与限制。",
            ],
            "logic": [
                "先检查这个 result 是否已经有直接子 revision，避免同一结果被重复审阅。",
                "Result-to-Claim 对每个主张给出 supported、partial 或 failed，写明测量基础、决定指标与不确定性。",
                "若 partial/failed，只有存在具体、可检验、预算可承受的下一动作才允许 requires_follow_up=true。",
                "Ablation Planner 只选择最能改变 Reviewer 信念的受控消融或诊断基线。",
                "PlanningAgent 根据 feedback 产生 revised_plan；保留已支持主张和控制变量，只改命名弱点要求的字段。",
                "生成 iteration_contract，记录 required_changes、required_metrics、required_comparisons 和 changed_fields。",
                "先保存 revision，再保存 parent=revision 的 refined plan，保证崩溃后仍能恢复完整决策。",
            ],
            "outputs": [
                "revision：verdict、supported/unsupported claims、revisions、next_action、evidence_links、overclaim_risks、iteration。",
                "refined plan：只有 requires_follow_up=true 且未达到上限时创建。",
                "反馈轮次保存在 RunRecord.feedback_iteration。",
            ],
            "validation": [
                "partial/failed 是有效完成状态，不强行改写为成功。",
                "代码实现修复必须冻结科学合同；科学修订才创建新 plan/bundle。",
                "报告导出前，最新 result 必须有属于它的 revision 且 requires_follow_up=false。",
                "当前 Engine 默认反馈上限可配置为 4；自动 Orchestrator 文件中仍有 2 轮常量，维护时应保持两处一致。",
            ],
        },
        {
            "index": "15",
            "title": "步骤 9：竞赛报告导出（report_export）",
            "subtitle": "只从经过验证的引用、真实实验和最终反馈生成可交付报告。",
            "agent": "Writer Agent",
            "skills": ["competition-report"],
            "model": "general（报告结构化生成）",
            "output": "report、下载文件",
            "inputs": [
                "所有经过筛选的 Artifact，但科学结论只能使用合格 evidence 和最新 audited result。",
                "最新 revision 必须已经结束 follow-up。",
            ],
            "logic": [
                "报告就绪门槛先检查至少一条 verified evidence。",
                "确认 result 属于最新 task/bundle 谱系，而不是旧迭代。",
                "竞赛模式要求 is_real_experiment=true。",
                "确认最新 revision.parent=result.id 且 requires_follow_up=false。",
                "Writer 生成问题、相关工作、假设、方法、实验、结果、局限、结论与引用。",
                "Competition Report 排除 uploaded-only 引用、失败或未审计指标及不受支持的主张。",
            ],
            "outputs": [
                "report Artifact、报告预览与下载接口。",
                "references、result_links、audit_summary，便于复核每条结论来源。",
            ],
            "validation": [
                "缺 evidence/result/revision、谱系不一致、仍需 follow-up 或非真实实验时拒绝导出。",
                "Reviewer 语义检查 verified references、audited results 和 unsupported claim absence。",
                "从实验步骤重跑时旧报告会失效，必须基于新结果重新导出。",
            ],
        },
    ]

    for step in steps:
        add_step_chapter(doc, step)
        if step["index"] == "10":
            add_figure(doc, diagrams["evidence"], "图 4  证据推理的双层审计与定向补证流程")
        if step["index"] == "13":
            add_figure(doc, diagrams["experiment"], "图 5  实验预检、正式执行、诊断和有界修复闭环")
        if step["index"] == "14":
            add_figure(doc, diagrams["lineage"], "图 6  Artifact 迭代谱系与当前实验重跑语义")

    section_title(doc, "16  停止、恢复、重新运行与历史保护")
    add_body(doc, "自动运行由 WorkflowOrchestrator 的后台线程驱动。它每轮从 Artifact 谱系计算 next_step，而不是只相信界面状态。用户点击停止后，Run 进入 stopping/paused，stop_requested=true；在途 Qwen 请求通过 cancel_run 中止，当前 Step 记为 interrupted。")
    add_table(
        doc,
        ["操作", "系统语义", "不会发生的事情"],
        [
            ("停止", "发出取消信号、保存当前状态、Run 进入 paused", "不会删除已完成 Artifact"),
            ("继续 Pipeline", "从谱系计算缺少的下一步骤", "不会无条件从步骤 1 重来"),
            ("重新运行普通步骤", "清理该步骤及其未锁定下游产物后重算", "不会删除锁定 Artifact"),
            ("重新运行本次实验", "保留当前迭代 plan/task/bundle，新增 attempt，仅使旧报告失效", "不会回退到初始实验"),
            ("服务重启恢复", "reconcile interrupted run，再由 Orchestrator 继续", "不会重复导入已匹配完成结果"),
        ],
        [1600, 4300, 3460],
        8.7,
    )
    add_body(doc, "“重新运行本次实验”会先记录 retry_mode=new_attempt_same_iteration，并写入 task_artifact_id、bundle_artifact_id 和 previous_result_artifact_id。旧 Result 仍保留；新 Result 包含累计 attempts。成功或失败后 force_new_attempt 标志都会清除。")
    add_callout(
        doc,
        "为什么之前会回到第一次实验",
        "旧实现把 feedback_revision 视为实验重跑的下游并删除，递归删除其子 plan、task 和 bundle，最终 latest task 退回 v1。现在实验重跑仅使 report_export 产物失效，不再破坏迭代谱系。",
        fill=LIGHT_GOLD,
        color=GOLD,
    )
    add_body(doc, "该设计与成熟工作流系统的原则一致：Argo 支持选择重试节点并决定是否重启成功节点；Temporal 以追加式 History 恢复状态。参考：https://github.com/argoproj/argo-workflows/discussions/7534 及 https://github.com/temporalio/temporal/blob/main/docs/architecture/history-service.md")

    section_title(doc, "17  使用方法简介（快速上手）")
    add_callout(doc, "最短路径", "设置数据和执行环境 → 新建课题 → 启动 Pipeline → 在证据/计划节点复核 → 观察实验 → 根据反馈继续或导出。", fill=LIGHT_GREEN, color=GREEN)
    doc.add_paragraph("A. 首次配置", style="Heading 2")
    for item in [
        "打开“项目设置”。选择实验环境：本地 GPU 或远程 SSH。",
        "本地 GPU：设置工作目录、Python 路径、CUDA 设备号；点击测试，确认目录、入口、Python 与 GPU 可用。",
        "数据集：选择“本地”时填写具体数据集目录；选择“在线/自动”时填写下载来源或使用系统目录中的可下载数据集。",
        "配置 Qwen API Key；确认 Provider 状态中 LLM、Literature、Experiment 均为 ready。",
        "不要把数据集缓存目录当作具体数据集目录。本地模式应直接指向本次实验唯一数据集。",
    ]:
        add_number(doc, item)
    if SCREENSHOT.is_file():
        add_figure(doc, SCREENSHOT, "图 7  本地数据集目录设置示例；目录会在生成假设前被检查并锁定", width=6.45)

    doc.add_paragraph("B. 创建并运行课题", style="Heading 2")
    for item in [
        "输入课题标题、研究问题、领域和约束。研究问题尽量写清目标、数据、基线、指标和资源限制。",
        "点击创建后启动 Pipeline。系统会依次运行九个步骤；长实验不需要浏览器一直保持前台。",
        "在“证据”区域检查引用是否 verified、是否存在 direct/indirect 支持，以及系统是否触发定向补证。",
        "在“假设”与“计划”区域检查主张、基线、唯一变化因素、seeds、指标、成功/失败阈值。",
        "实验运行时查看 variant、seed、epoch、loss 和日志；不要把实时 loss 当最终指标。",
        "若反馈为 partial/failed，阅读 required_revision 和 next_action；系统会在条件满足时生成下一轮计划。",
        "当 revision 不再要求 follow-up 且结果通过审计后，导出竞赛报告。",
    ]:
        add_number(doc, item)
    doc.add_paragraph("C. 停止与重跑", style="Heading 2")
    for item in [
        "需要中断时点击停止，等待状态变为“已暂停”。",
        "要继续整条流程，使用继续/启动 Pipeline，让系统根据谱系选择下一步。",
        "只想重做当前实验时，使用“重新运行本次实验”；它会复用当前迭代 Bundle 并新增 attempt。",
        "若反馈已经产生新计划但 Bundle 尚未生成，应继续 Pipeline 生成新 task/bundle，而不是重跑旧实验。",
        "修改数据集、科学假设或计划属于新科学合同；不应伪装成代码修复。",
    ]:
        add_bullet(doc, item)

    section_title(doc, "18  常见问题与排查方法")
    add_table(
        doc,
        ["现象", "含义/原因", "处理方法"],
        [
            ("证据推理失败", "断网、外部检索失败、无候选通过证据硬门槛、模型 JSON 形状多次不合格", "恢复网络；查看 targeted retrieval、warnings、candidate issues；不要强制 GO"),
            ("本地数据集检查失败", "目录缺失、为空、指向缓存根而非具体数据集、内容变化导致指纹不一致", "重新选择具体目录；确认文件权限与内容稳定"),
            ("Bundle 生成失败", "语法、依赖、指标来源、参数协议或数据集契约不合格", "查看 Supervisor rejection；让代码模型按错误修复"),
            ("实验启动即失败", "CUDA、依赖、API、数据 Schema、参数或环境问题", "查看 diagnosis.category/error_code 和预检日志"),
            ("训练长时间无输出", "代码未输出 JSONL epoch_end、进程静默但仍存活、或实际卡住", "查看进程状态和日志时间戳；不要仅凭耗时判失败"),
            ("结果有数值但不能导出", "audit 未通过、非真实 GPU 实验、谱系错误、未完成反馈或仍需 follow-up", "先处理 audit issues 和 revision"),
            ("重跑后 attempt 没增加", "恢复了完全匹配的已完成结果，或没有设置强制新 attempt", "使用“重新运行本次实验”，检查 retry 事件"),
            ("继续后生成新 Bundle", "最新 Artifact 是 refined plan，而不是可复用的旧 Bundle", "这是科学迭代的正确行为"),
        ],
        [1700, 4100, 3560],
        8.2,
    )
    add_callout(
        doc,
        "排查优先级",
        "先看 Run 状态和 current_step → 再看该 Step.error → 再看最新 Event 的 actor/message/output_summary → 再检查 Artifact 谱系 → 最后查看实验日志与 Provider 配置。",
    )

    section_title(doc, "19  安全、可复现与科研诚信边界")
    for item in [
        "不能把未验证引用、上传文件名或模型记忆当成正式证据。",
        "不能把 stdout 文本解析出的数字当最终实验指标；指标必须来自声明的 result JSON。",
        "不能把训练进程退出码 0 当科学结论；还需 analysis、audit 和 result-to-claim。",
        "不能在代码修复时修改 dataset、seeds、metrics、GPU requirement 或科学计划。",
        "不能自动安装依赖、修改 CUDA/驱动、编辑凭据或删除未知目录。",
        "不能把同一 result 重复审阅成多个反馈轮次。",
        "不能把 partial/failed 改写为成功；负结果和不确定结果必须保留。",
        "报告导出只能使用 exportable verified citations 与 audited real experiment。",
    ]:
        add_bullet(doc, item)
    add_body(doc, "可复现性由多层共同保证：数据指纹、Task/Bundle/Result ID、父 Artifact、源码哈希、参数与 seeds、环境摘要、append-only attempts、事件时间线、模型路由和 Skill 指令哈希。")

    section_title(doc, "20  当前配置快照与实现注意事项")
    add_table(
        doc,
        ["类别", "当前状态", "说明"],
        [
            ("LLM", "qwen，ready", "推理实测 qwen3.7-max；general 实测 qwen-max"),
            ("Literature", "arxiv_semantic_scholar，ready", "并结合 Research Wiki 与本地文献"),
            ("Experiment", "local_gpu，ready", "CUDA_VISIBLE_DEVICES=0"),
            ("Dataset", "local: datasets/IPIX", "启动前检查并绑定目录内容指纹"),
            ("下载重试", "5", "仅用于允许下载的数据集模式"),
            ("实验超时", "0", "0 表示不按固定时长强制超时"),
            ("反馈上限", "Engine 默认 4；Orchestrator 常量 2", "建议后续统一为单一配置来源"),
            ("测试状态", "387 passed，2 skipped", "包含当前迭代重跑谱系回归测试"),
        ],
        [1700, 2550, 5110],
        8.7,
    )
    add_callout(
        doc,
        "敏感信息",
        "本报告不包含 Qwen API Key、SSH 私钥、密码或其他凭据。Provider 页面只应显示是否已配置，不应显示明文。",
        fill=LIGHT_GOLD,
        color=GOLD,
    )

    section_title(doc, "附录 A  API 与界面操作映射")
    add_table(
        doc,
        ["界面动作", "HTTP 方法与路径", "效果"],
        [
            ("创建课题", "POST /api/runs", "创建 RunRecord 与九个 StepRecord"),
            ("读取课题", "GET /api/runs/{run_id}", "获取完整状态、Artifact 与 Event"),
            ("运行单步", "POST /api/runs/{run_id}/steps/{step}/run", "同步执行一个步骤"),
            ("启动 Pipeline", "POST /api/runs/{run_id}/pipeline/start", "后台自动推进"),
            ("停止 Pipeline", "POST /api/runs/{run_id}/pipeline/stop", "取消在途请求并暂停"),
            ("从步骤重跑", "POST /api/runs/{run_id}/steps/{step}/rerun-from", "按重跑规则清理/保留谱系后执行"),
            ("实验进度", "GET /api/runs/{run_id}/experiment-progress", "读取进程、JSONL progress 与结果状态"),
            ("终止实验", "POST /api/runs/{run_id}/experiments/{experiment_id}/terminate", "终止指定实验进程"),
            ("Provider 状态", "GET /api/settings/providers", "检查 LLM/文献/实验是否 ready"),
            ("实验设置", "GET/PUT /api/settings/experiment", "读取或更新本地/远程/数据集设置"),
            ("测试设置", "POST /api/settings/experiment/test", "检查目录、Python、入口和执行环境"),
            ("导出报告", "GET /api/runs/{run_id}/report/download", "下载最终报告"),
        ],
        [1700, 4200, 3460],
        8.3,
    )

    section_title(doc, "附录 B  主要实现文件索引")
    add_table(
        doc,
        ["文件", "说明"],
        [
            ("backend/app/workflow/engine.py", "九步执行、验证循环、诊断修复、反馈、重跑与报告门槛"),
            ("backend/app/workflow/orchestrator.py", "后台自动推进、停止恢复和 Artifact 谱系 next_step"),
            ("backend/app/workflow/skills.py", "步骤 → Agent/Skill 静态路由与条件 Skill"),
            ("backend/app/workflow/skill_runtime.py", "Skill 加载、指令预算、工具交集授权和审计哈希"),
            ("backend/app/workflow/evidence_audit.py", "证据注册表、候选审计、定向补证查询"),
            ("backend/app/workflow/knowledge.py", "Wiki/本地/外部检索与 verified evidence 合并"),
            ("backend/app/providers/llm.py", "Qwen 模型路由、重试、取消、JSON 解析和调用元数据"),
            ("backend/app/providers/experiment.py", "本地 GPU、远程 SSH、结果恢复、运行与数据可用性"),
            ("backend/app/workflow/experiment_code.py", "Bundle 规范化、源码与 manifest 验证"),
            ("backend/app/models/*.py", "Run、Step、Artifact、Event、ExperimentBundle 等数据模型"),
            ("frontend/src/App.tsx", "课题创建、自动运行、单步、停止与重跑交互"),
            ("frontend/src/components/*.tsx", "时间线、证据表、实验面板、设置与报告预览"),
            ("skills/*/SKILL.md", "每个科研环节的执行协议、工具、边界和输出合同"),
        ],
        [3600, 5760],
        8.6,
    )

    section_title(doc, "附录 C  术语表")
    add_table(
        doc,
        ["术语", "定义"],
        [
            ("Run", "一次完整科研流程的持久化运行实例。"),
            ("Step", "九步状态机中的一个可执行状态转换。"),
            ("Artifact", "步骤产生的结构化科研产物；可锁定并通过 parent_artifact_id 形成谱系。"),
            ("Event", "发生过的调用、决定、错误、修订和模型/工具元数据。"),
            ("Agent", "负责某类领域任务的执行组件。"),
            ("Skill", "注入 Agent 的操作协议、边界、工具声明和输出合同，不是独立模型。"),
            ("Provider", "外部能力适配器，如 Qwen、文献检索、本地 GPU 或 SSH。"),
            ("Bundle", "可执行实验源码、依赖和 manifest 的不可变组合。"),
            ("Attempt", "同一实验 Bundle 的一次实际执行记录，采用追加式编号。"),
            ("Audit", "对来源、源码、命令、环境、指标和谱系的独立完整性检查。"),
            ("Revision", "基于某个具体 Result 的科学结论与下一步修订决定。"),
            ("MDE", "Minimum Decisive Experiment，能够改变决策的最小判别性实验。"),
        ],
        [1900, 7460],
        8.8,
    )

    section_title(doc, "结语")
    add_body(doc, "这套系统的核心价值不是“自动生成一篇看起来完整的报告”，而是把科研决策拆成可以验证、重做和追责的步骤：证据必须有来源，假设必须能证伪，实验必须有父计划和不可变 Bundle，结果必须经过审计，迭代必须针对具体弱点，报告必须只使用最终可支持的主张。")
    add_callout(
        doc,
        "推荐日常使用习惯",
        "每到证据推理、研究计划、实验审计和反馈修订四个节点，人工快速复核一次；其余步骤可以自动推进。遇到错误先停止并看谱系，不要连续点击重跑。",
        fill=LIGHT_GREEN,
        color=GREEN,
    )

    # Prevent table rows from splitting and apply widow control.
    for table in doc.tables:
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            cant_split = OxmlElement("w:cantSplit")
            tr_pr.append(cant_split)
    for p in doc.paragraphs:
        p_pr = p._p.get_or_add_pPr()
        widow = OxmlElement("w:widowControl")
        p_pr.append(widow)

    doc.core_properties.title = "AI 科研实验系统：执行逻辑、技术架构与使用指南"
    doc.core_properties.subject = "九步科研工作流、模型路由、Skill、实验迭代与操作手册"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "AI Scientist, Qwen, Evidence Reasoning, Experiment Workflow, Skill Runtime"
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_report()
