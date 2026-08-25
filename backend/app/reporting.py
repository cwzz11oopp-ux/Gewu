from __future__ import annotations

from io import BytesIO
from html import escape
import json
import re
from typing import Any
from urllib.parse import quote
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from backend.app.report_visualization import FigureSpec, ReportSpec, render_figure_png


SECTION_ALIASES = (
    ("执行摘要", ("Executive Summary", "执行摘要")),
    ("一、待研究问题", ("Problem Statement", "待研究问题")),
    ("二、解决思路与推导依据", ("Rationale", "解决思路")),
    ("三、必要的技术手段", ("Technical Details", "必要的技术手段")),
    ("四、数据集", ("Datasets", "数据集")),
    ("五、Source：假设推演依据", ("Source",)),
    ("六、Target：拟验证的数据特征", ("Target",)),
    ("七、方法论", ("Methods", "方法论")),
    ("八、实验设计", ("Experiments", "实验设计")),
    ("九、实验迭代过程", ("Iteration Summary", "实验迭代过程")),
    ("十、实验结果", ("Results", "实验结果")),
    ("十一、研究结论与后续判断", ("Research Conclusion", "研究结论")),
    ("十二、局限性", ("Limitations", "局限性")),
    ("十三、复现说明", ("Reproducibility", "复现信息")),
    ("十四、参考文献", ("References", "参考文献")),
)

BLUE = RGBColor(31, 78, 121)
DARK = RGBColor(31, 41, 55)
MUTED = RGBColor(100, 116, 139)
LIGHT_FILL = "F2F4F7"
TABLE_INDENT_DXA = 120
TABLE_CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

FIGURE_SECTION = {
    "model_structure": "method",
    "method_pipeline": "design",
    "control_variables": "design",
    "workflow_timeline": "iteration",
    "training_curve": "iteration",
    "main_comparison": "results",
    "seed_comparison": "results",
    "seed_delta": "results",
}


def render_report_html(report: dict, *, run_id: str, run_title: str) -> str:
    """Legacy preview helper kept for callers; the export API no longer ships HTML."""
    title = str(_pick(report, "Paper Title", "标题") or run_title or "科学假设与研究报告")
    status = report.get("Report Status") if isinstance(report.get("Report Status"), dict) else {}
    experiment_label = "真实实验" if status.get("real_experiment") is True else "未通过真实实验审计"
    sections = []
    for heading, aliases in SECTION_ALIASES:
        value = _pick(report, *aliases)
        if value in (None, "", [], {}):
            continue
        sections.append(f"<section><h2>{escape(heading)}</h2><pre>{escape(_short_text(value, 6000))}</pre></section>")
    return (
        "<!doctype html><html lang=\"zh-CN\"><head><meta charset=\"utf-8\">"
        "<style>body{font:15px/1.7 sans-serif;max-width:900px;margin:auto;padding:40px}"
        "h2{color:#1f4e79}pre{white-space:pre-wrap}@media print{body{padding:0}}</style>"
        f"<title>{escape(title)}</title></head><body><h1>{escape(title)}</h1>"
        f"<p>{experiment_label} · {escape(run_id)}</p>{''.join(sections)}</body></html>"
    )


def build_report_docx(report: dict, *, run_id: str, run_title: str) -> bytes:
    document = Document()
    _configure_document(document)
    title = str(
        _pick(report, "Report Title", "Paper Title", "标题")
        or run_title
        or "科学假设与研究报告"
    )
    subtitle = "科学假设与研究报告"
    _add_title_block(document, title, subtitle, run_id)

    abstract = _pick(report, "Paper Abstract", "摘要")
    if abstract:
        document.add_heading("摘要", level=1)
        _add_prose(document, abstract)
        keywords = report.get("Keywords") or []
        if keywords:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(12)
            paragraph.add_run("关键词：").bold = True
            paragraph.add_run("；".join(str(item) for item in keywords[:6]))

    narrative_sections = report.get("Narrative Sections")
    if isinstance(narrative_sections, list) and narrative_sections:
        _add_narrative_sections(document, narrative_sections, report)
        references = report.get("References")
        if references:
            document.add_heading("参考文献", level=1)
            _add_references(document, references)
    else:
        _add_report_figures(document, report)
        for heading, aliases in SECTION_ALIASES:
            value = _pick(report, *aliases)
            if value in (None, "", [], {}):
                continue
            document.add_heading(heading, level=1)
            if heading == "四、数据集":
                _add_dataset(document, value)
            elif heading == "八、实验设计":
                _add_experiment_design(document, value)
            elif heading == "九、实验迭代过程":
                _add_iteration_summary(document, value)
            elif heading == "十、实验结果":
                _add_results(document, value)
            elif heading == "十三、复现说明":
                _add_reproducibility(document, value)
            elif heading == "十四、参考文献":
                _add_references(document, value)
            else:
                _add_human_value(document, value)

    _add_footer(document, run_id)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_report_zip(report: dict, *, run_id: str, run_title: str, artifacts: list) -> bytes:
    """Create the deliberately small primary package requested by the user.

    It contains only the Word report, experiment source code, and experiment
    results. Internal artifact snapshots, HTML, Markdown, literature payloads,
    and package manifests are intentionally excluded.
    """
    stream = BytesIO()
    with ZipFile(stream, "w", ZIP_DEFLATED) as archive:
        archive.writestr(
            "研究报告/科学假设与研究报告.docx",
            build_report_docx(report, run_id=run_id, run_title=run_title),
        )
        _write_experiment_code(archive, artifacts)
        _write_experiment_results(archive, artifacts)
    return stream.getvalue()


def _add_report_figures(document: Document, report: dict) -> None:
    """Legacy reports place all grounded figures in one block."""
    for number, figure in _report_figures(report):
        _add_figure(document, figure, number)
    raw_spec = report.get("Report Spec")
    if not isinstance(raw_spec, dict):
        return
    spec = ReportSpec.model_validate(raw_spec)
    for omitted in spec.omitted_figures:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(f"未生成 {omitted.figure_id}：{omitted.reason}")
        _set_run_font(run, "宋体", 9, MUTED)


def _report_figures(report: dict) -> list[tuple[int, FigureSpec]]:
    raw_spec = report.get("Report Spec")
    if not isinstance(raw_spec, dict):
        return []
    spec = ReportSpec.model_validate(raw_spec)
    return list(enumerate(spec.figures, 1))


def _add_section_figures(document: Document, report: dict, section_id: str) -> None:
    for number, figure in _report_figures(report):
        if FIGURE_SECTION.get(figure.figure_id) == section_id:
            _add_figure(document, figure, number)


def _add_figure(document: Document, figure: FigureSpec, number: int) -> None:
    png = render_figure_png(figure)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(BytesIO(png), width=Inches(6.45))

    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(10)
    run = caption.add_run(f"图 {number}  {figure.title}。{figure.caption}")
    _set_run_font(run, "宋体", 9, DARK)


def build_experiment_package(*, run_id: str, artifacts: list) -> bytes:
    """Package only experiment assets that actually exist in persisted artifacts."""
    stream = BytesIO()
    with ZipFile(stream, "w", ZIP_DEFLATED) as archive:
        _write_experiment_code(archive, artifacts)
        _write_experiment_results(archive, artifacts)
        for artifact in artifacts:
            if getattr(artifact, "type", "") != "experiment_result":
                continue
            version = getattr(artifact, "version", 1)
            content = getattr(artifact, "content", {})
            for key, folder in (("parameters", "configs"), ("environment", "environment"), ("metrics", "metrics"), ("attempts", "logs"), ("audit", "artifacts")):
                value = content.get(key)
                if value not in (None, "", [], {}):
                    archive.writestr(
                        f"{folder}/第{version}轮-{key}.json",
                        json.dumps(value, ensure_ascii=False, indent=2, default=str).encode("utf-8"),
                    )
    return stream.getvalue()


def attachment_headers(filename: str) -> dict[str, str]:
    fallback = filename if filename.isascii() else "research-output.zip"
    return {
        "Content-Disposition": (
            f'attachment; filename="{fallback}"; filename*=UTF-8\'\'{quote(filename)}'
        )
    }


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, before, after in (
        ("Title", 28, 0, 12),
        ("Heading 1", 16, 18, 10),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 12, 8, 4),
    ):
        style = styles[name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLUE if name != "Title" else DARK
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def _add_title_block(document: Document, title: str, subtitle: str, run_id: str) -> None:
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(8)
    run = kicker.add_run("格物 · 研究报告")
    _set_run_font(run, "微软雅黑", 11, BLUE, True)

    paragraph = document.add_paragraph(style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.add_run(title)

    description = document.add_paragraph()
    description.alignment = WD_ALIGN_PARAGRAPH.CENTER
    description.paragraph_format.space_after = Pt(8)
    run = description.add_run(subtitle)
    _set_run_font(run, "宋体", 11, DARK)

    metadata = document.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.paragraph_format.space_after = Pt(16)
    run = metadata.add_run(f"运行编号：{run_id}")
    _set_run_font(run, "宋体", 9, MUTED)


def _add_human_value(document: Document, value: Any) -> None:
    if isinstance(value, str):
        _add_prose(document, value)
        return
    if isinstance(value, list):
        for item in value:
            if isinstance(item, dict):
                _add_compact_record(document, item)
            else:
                _add_bullet(document, str(item))
        return
    if isinstance(value, dict):
        for key, item in _curated_items(value):
            if isinstance(item, (dict, list)):
                document.add_heading(_human_label(key), level=2)
                _add_human_value(document, item)
            else:
                paragraph = document.add_paragraph()
                label = paragraph.add_run(f"{_human_label(key)}：")
                label.bold = True
                paragraph.add_run(str(item))
        return
    _add_prose(document, str(value))


def _add_narrative_sections(
    document: Document,
    sections: list,
    report: dict,
) -> None:
    """Render the Qwen-authored narrative and only a few supporting data tables."""
    for section in sections:
        if not isinstance(section, dict):
            continue
        title = str(section.get("title") or "").strip()
        if not title:
            continue
        document.add_heading(title, level=1)
        for paragraph in section.get("paragraphs") or []:
            _add_prose(document, paragraph)
        for subsection in section.get("subsections") or []:
            if not isinstance(subsection, dict):
                continue
            subtitle = str(subsection.get("title") or "").strip()
            if subtitle:
                document.add_heading(subtitle, level=2)
            for paragraph in subsection.get("paragraphs") or []:
                _add_prose(document, paragraph)

        section_id = str(section.get("id") or "")
        if section_id == "design":
            _add_design_support_table(document, report.get("Experiments"), table_number=1)
        elif section_id == "results":
            _add_result_support_table(document, report.get("Results"), table_number=2)
        elif section_id == "conclusion":
            _add_reproducibility_support_table(document, report.get("Reproducibility"), table_number=4)
        _add_section_figures(document, report, section_id)


def _add_design_support_table(document: Document, value: Any, *, table_number: int) -> None:
    if not isinstance(value, dict):
        return
    parameters = value.get("parameters") if isinstance(value.get("parameters"), dict) else {}
    rows = [(_human_label(key), _short_text(item, 240)) for key, item in _public_parameters(parameters)]
    seeds = value.get("seeds") or []
    if seeds:
        rows.append(("随机种子", "、".join(str(item) for item in seeds)))
    if rows:
        _add_table_caption(document, table_number, "主要实验参数与控制变量")
        _add_label_table(document, rows)


def _add_result_support_table(document: Document, value: Any, *, table_number: int) -> None:
    if not isinstance(value, dict):
        return
    summary = value.get("metric_summary") if isinstance(value.get("metric_summary"), dict) else {}
    pairs = _preferred_metric_pairs(_resolve_metric_pairs(summary, value_kind="summary"))
    accuracy_pairs = [pair for pair in pairs if "accuracy" in _normalized_metric_name(pair[0])]
    shown_pairs = accuracy_pairs or pairs[:3]
    rows = []
    for metric_name, baseline_key, experiment_key in shown_pairs:
        baseline = summary[baseline_key]
        experiment = summary[experiment_key]
        baseline_mean = baseline.get("mean")
        experiment_mean = experiment.get("mean")
        if not isinstance(baseline_mean, (int, float)) or not isinstance(experiment_mean, (int, float)):
            continue
        rows.append(
            (
                _human_label(metric_name),
                _format_mean_std(metric_name, baseline),
                _format_mean_std(metric_name, experiment),
                _format_metric_delta(metric_name, float(experiment_mean) - float(baseline_mean)),
            )
        )
    if rows:
        _add_table_caption(document, table_number, "基线模型与实验模型的主要结果")
        _add_table(document, ("指标", "基线模型", "实验模型", "差值"), rows, (2.2, 1.5, 1.5, 1.3))
    else:
        metrics = value.get("metrics") if isinstance(value.get("metrics"), dict) else {}
        if metrics:
            _add_table_caption(document, table_number, "最终实验指标")
            fallback_rows = [
                (_human_label(key), _format_number(item))
                for key, item in _deduplicated_metrics(metrics)
            ]
            _add_table(document, ("指标", "结果"), fallback_rows, (3.75, 2.75))

    seeds = value.get("seed_results") if isinstance(value.get("seed_results"), list) else []
    seed_rows = []
    seed_metric_maps = [
        item.get("metrics") if isinstance(item, dict) and isinstance(item.get("metrics"), dict) else {}
        for item in seeds
    ]
    seed_pairs = (
        _preferred_metric_pairs(_resolve_metric_pairs(seed_metric_maps[0], value_kind="number"))
        if seed_metric_maps
        else []
    )
    accuracy_seed_pairs = [pair for pair in seed_pairs if "accuracy" in _normalized_metric_name(pair[0])]
    selected_seed_pair = (accuracy_seed_pairs or seed_pairs[:1])[:1]
    if selected_seed_pair:
        metric_name, baseline_key, experiment_key = selected_seed_pair[0]
        for index, (item, metrics) in enumerate(zip(seeds, seed_metric_maps)):
            baseline = metrics.get(baseline_key)
            experiment = metrics.get(experiment_key)
            if not isinstance(baseline, (int, float)) or not isinstance(experiment, (int, float)):
                seed_rows = []
                break
            seed_rows.append(
                (
                    str(item.get("seed") or index + 1),
                    f"{float(baseline) * 100:.2f}%",
                    f"{float(experiment) * 100:.2f}%",
                    f"{(float(experiment) - float(baseline)) * 100:+.2f} pp",
                )
            )
    if seed_rows:
        _add_table_caption(document, table_number + 1, f"逐随机种子的{_human_label(metric_name)}配对结果")
        _add_table(
            document,
            ("种子", "基线方案", "实验方案", "差值"),
            seed_rows,
            (1.35, 1.7, 1.7, 1.75),
        )


def _add_reproducibility_support_table(document: Document, value: Any, *, table_number: int) -> None:
    if not isinstance(value, dict):
        return
    rows = []
    if value.get("experiment_id"):
        rows.append(("实验编号", str(value["experiment_id"])))
    if value.get("seeds"):
        rows.append(("随机种子", "、".join(str(item) for item in value["seeds"])))
    if isinstance(value.get("parameters"), dict) and value["parameters"]:
        public_parameters = dict(_public_parameters(value["parameters"]))
        if public_parameters:
            rows.append(("关键参数", _short_text(public_parameters, 700)))
    environment = value.get("environment")
    if isinstance(environment, dict):
        public_environment = {
            key: environment[key]
            for key in ("python", "torch", "cuda", "device")
            if environment.get(key) not in (None, "")
        }
        if public_environment:
            rows.append(("运行环境", _short_text(public_environment, 400)))
    if rows:
        _add_table_caption(document, table_number, "关键复现信息")
        _add_label_table(document, rows)


def _add_table_caption(document: Document, number: int, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(f"表 {number}  {title}")
    _set_run_font(run, "宋体", 9, DARK)


def _format_mean_std(metric_name: str, value: dict[str, Any]) -> str:
    mean = value.get("mean")
    std = value.get("std")
    if not isinstance(mean, (int, float)):
        return "—"
    if "Parameters" in metric_name:
        return f"{float(mean):,.0f}" + (f" ± {float(std):,.0f}" if isinstance(std, (int, float)) else "")
    return f"{float(mean) * 100:.2f}%" + (f" ± {float(std) * 100:.2f}%" if isinstance(std, (int, float)) else "")


def _format_metric_delta(metric_name: str, value: float) -> str:
    if "Parameters" in metric_name:
        return f"{value:+,.0f}"
    return f"{value * 100:+.2f} pp"


def _add_prose(document: Document, value: Any) -> None:
    text = str(value).strip()
    if not text:
        return
    for block in (part.strip() for part in text.split("\n") if part.strip()):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Inches(0.28)
        paragraph.add_run(block)


def _add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.add_run(text.strip())


def _add_compact_record(document: Document, value: dict) -> None:
    summary = []
    for key, item in _curated_items(value):
        if item in (None, "", [], {}):
            continue
        if isinstance(item, (str, int, float, bool)):
            summary.append(f"{_human_label(key)}：{item}")
    if summary:
        _add_bullet(document, "；".join(summary))


def _add_dataset(document: Document, value: Any) -> None:
    if not isinstance(value, dict):
        _add_human_value(document, value)
        return
    allowed = (
        "name", "source", "source_type", "split", "preprocessing",
        "limitations", "availability", "license", "description",
    )
    rows = []
    for key in allowed:
        item = value.get(key)
        if item not in (None, "", [], {}):
            rows.append((_human_label(key), _short_text(item, 900)))
    card = value.get("card")
    if isinstance(card, dict):
        for key in ("name", "source_type", "availability", "license"):
            if card.get(key) not in (None, "") and all(label != _human_label(key) for label, _ in rows):
                rows.append((_human_label(key), _short_text(card[key], 300)))
    if rows:
        _add_label_table(document, rows)
    else:
        _add_prose(document, "当前产物仅记录了数据集名称，未记录更多可公开说明的信息。")


def _add_experiment_design(document: Document, value: Any) -> None:
    if not isinstance(value, dict):
        _add_human_value(document, value)
        return
    comparisons = value.get("comparisons") or value.get("对比结果")
    if comparisons:
        document.add_heading("对照方案", level=2)
        _add_human_value(document, comparisons)
    parameters = value.get("parameters") or value.get("参数")
    if isinstance(parameters, dict) and parameters:
        document.add_heading("关键参数", level=2)
        _add_label_table(
            document,
            [(_human_label(key), _short_text(item, 250)) for key, item in _public_parameters(parameters)],
        )
    seeds = value.get("seeds") or value.get("随机种子")
    if seeds:
        paragraph = document.add_paragraph()
        paragraph.add_run("随机种子：").bold = True
        paragraph.add_run("、".join(str(item) for item in seeds))
    evaluations = value.get("evaluations") or value.get("指标")
    if evaluations:
        document.add_heading("评估指标与判定标准", level=2)
        _add_human_value(document, evaluations)
    procedure = value.get("procedure")
    if isinstance(procedure, dict):
        procedure = procedure.get("steps")
    if procedure:
        document.add_heading("实验步骤", level=2)
        for index, item in enumerate(procedure, 1):
            paragraph = document.add_paragraph(style="List Number")
            paragraph.add_run(str(item))


def _add_iteration_summary(document: Document, value: Any) -> None:
    if not isinstance(value, dict):
        _add_human_value(document, value)
        return
    rounds = value.get("rounds") or value.get("各轮记录") or []
    if not isinstance(rounds, list) or not rounds:
        _add_prose(document, "当前报告未记录逐轮实验信息。")
        return
    rows = []
    for index, item in enumerate(rounds, 1):
        if not isinstance(item, dict):
            continue
        rows.append(
            (
                str(item.get("round") or index),
                str(item.get("experiment_id") or item.get("experiment id") or ""),
                _short_text(item.get("metrics") or item.get("指标") or "", 280),
                str(item.get("feedback_verdict") or item.get("feedback verdict") or item.get("status") or ""),
                _short_text(item.get("required_revision") or item.get("required revision") or "", 320),
            )
        )
    _add_table(document, ("轮次", "实验", "关键结果", "判断", "本轮处理"), rows, (0.5, 1.0, 2.05, 0.75, 2.2))


def _add_results(document: Document, value: Any) -> None:
    if not isinstance(value, dict):
        _add_human_value(document, value)
        return
    metrics = value.get("metrics") if isinstance(value.get("metrics"), dict) else {}
    if metrics:
        rows = [(_human_label(key), _format_number(item)) for key, item in metrics.items()]
        _add_table(document, ("指标", "结果"), rows, (3.75, 2.75))
    analysis = value.get("analysis")
    if isinstance(analysis, dict):
        for key in ("summary", "conclusion", "interpretation", "limitations"):
            if analysis.get(key):
                document.add_heading(_human_label(key), level=2)
                _add_human_value(document, analysis[key])
    verdict = value.get("verdict") or value.get("status")
    if verdict:
        paragraph = document.add_paragraph()
        paragraph.add_run("实验状态：").bold = True
        paragraph.add_run(str(verdict))


def _add_reproducibility(document: Document, value: Any) -> None:
    if not isinstance(value, dict):
        _add_human_value(document, value)
        return
    allowed = ("experiment_id", "provider", "parameters", "seeds", "environment")
    rows = []
    for key in allowed:
        item = value.get(key)
        if item in (None, "", [], {}):
            continue
        if key == "environment" and isinstance(item, dict):
            item = {k: item[k] for k in ("python", "torch", "cuda", "device") if k in item}
        if key == "parameters" and isinstance(item, dict):
            item = dict(_public_parameters(item))
        rows.append((_human_label(key), _short_text(item, 600)))
    _add_label_table(document, rows)


def _add_references(document: Document, value: Any) -> None:
    if not isinstance(value, list):
        _add_human_value(document, value)
        return
    for index, reference in enumerate(value[:15], 1):
        if not isinstance(reference, dict):
            document.add_paragraph(f"[{index}] {reference}")
            continue
        authors = reference.get("authors") or []
        if isinstance(authors, list):
            authors = ", ".join(str(item) for item in authors[:6])
            if len(reference.get("authors") or []) > 6:
                authors += ", et al."
        identifiers = reference.get("identifiers") if isinstance(reference.get("identifiers"), dict) else {}
        identifier = identifiers.get("doi") or identifiers.get("arxiv") or reference.get("url") or ""
        citation = " ".join(
            str(item).strip()
            for item in (authors, f"({reference.get('year')})." if reference.get("year") else "", reference.get("title"), identifier)
            if item
        )
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.28)
        paragraph.paragraph_format.first_line_indent = Inches(-0.28)
        paragraph.add_run(f"[{index}] {citation}")


def _add_label_table(document: Document, rows: list[tuple[str, str]]) -> None:
    _add_table(document, ("项目", "说明"), rows, (1.5, 5.0))


def _add_table(document: Document, headers: tuple[str, ...], rows: list[tuple], widths: tuple[float, ...]) -> None:
    if not rows:
        return
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.style = "Table Grid"
    widths_dxa = tuple(int(round(width * 1440)) for width in widths)
    _set_table_geometry(table, widths_dxa)
    for index, (header, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[index]
        cell.width = Inches(width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _shade_cell(cell, LIGHT_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, (item, width) in enumerate(zip(row, widths)):
            cells[index].width = Inches(width)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.add_run(str(item))
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def _set_table_geometry(table, widths_dxa: tuple[int, ...]) -> None:
    """Apply fixed geometry so Word and LibreOffice wrap tables identically."""
    table_properties = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout", "w:tblCellMar"):
        existing = table_properties.find(qn(tag))
        if existing is not None:
            table_properties.remove(existing)

    table_width = OxmlElement("w:tblW")
    table_width.set(qn("w:w"), str(sum(widths_dxa)))
    table_width.set(qn("w:type"), "dxa")
    table_properties.append(table_width)

    table_indent = OxmlElement("w:tblInd")
    table_indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    table_indent.set(qn("w:type"), "dxa")
    table_properties.append(table_indent)

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table_properties.append(layout)

    margins = OxmlElement("w:tblCellMar")
    for side, value in TABLE_CELL_MARGIN_DXA.items():
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    table_properties.append(margins)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)

    for row in table.rows:
        for cell, value in zip(row.cells, widths_dxa):
            properties = cell._tc.get_or_add_tcPr()
            width = properties.find(qn("w:tcW"))
            if width is None:
                width = OxmlElement("w:tcW")
                properties.append(width)
            width.set(qn("w:w"), str(value))
            width.set(qn("w:type"), "dxa")


def _add_footer(document: Document, run_id: str) -> None:
    section = document.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header_paragraph.add_run("格物 · 科学研究报告")
    _set_run_font(header_run, "宋体", 8.5, MUTED)
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(f"{run_id}  ·  ")
    _set_run_font(run, "宋体", 8.5, MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def _write_experiment_code(archive: ZipFile, artifacts: list) -> None:
    for artifact in artifacts:
        if getattr(artifact, "type", "") != "experiment_bundle":
            continue
        content = getattr(artifact, "content", {})
        version = getattr(artifact, "version", 1)
        for item in content.get("files") or []:
            if not isinstance(item, dict) or not item.get("path") or "content" not in item:
                continue
            safe_path = str(item["path"]).replace("\\", "/").lstrip("/")
            if ".." not in safe_path.split("/"):
                archive.writestr(f"实验代码/第{version}版/{safe_path}", str(item["content"]).encode("utf-8"))
        requirements = content.get("requirements") or []
        if requirements:
            archive.writestr(
                f"实验代码/第{version}版/requirements.txt",
                "\n".join(str(item) for item in requirements).encode("utf-8"),
            )


def _write_experiment_results(archive: ZipFile, artifacts: list) -> None:
    for artifact in artifacts:
        if getattr(artifact, "type", "") != "experiment_result":
            continue
        version = getattr(artifact, "version", 1)
        content = getattr(artifact, "content", {})
        archive.writestr(
            f"实验结果/第{version}轮实验结果.json",
            json.dumps(content, ensure_ascii=False, indent=2, default=str).encode("utf-8"),
        )


def _curated_items(value: dict) -> list[tuple[str, Any]]:
    hidden = {
        "files", "sample_sha256", "sha256", "content_fingerprint", "artifact_id",
        "parent_artifact_id", "attempts", "audit", "command", "workdir", "log_path",
        "result_path", "bundle", "provider_mode", "fallback_used", "fallback_reason",
    }
    return [(str(key), item) for key, item in value.items() if str(key) not in hidden]


def _public_parameters(value: dict[str, Any]) -> list[tuple[str, Any]]:
    """Return reader-facing scalar controls and omit planning metadata containers."""
    hidden = {"additional_sections", "training_budget_rationale"}
    return [
        (str(key), item)
        for key, item in value.items()
        if str(key) not in hidden
        and item not in (None, "")
        and isinstance(item, (str, int, float, bool))
    ]


def _resolve_metric_pairs(
    metrics: dict[str, Any],
    *,
    value_kind: str,
) -> list[tuple[str, str, str]]:
    def valid(item: Any) -> bool:
        if value_kind == "summary":
            return isinstance(item, dict) and isinstance(item.get("mean"), (int, float))
        return isinstance(item, (int, float)) and not isinstance(item, bool)

    pairs: list[tuple[str, str, str]] = []
    for key, baseline in metrics.items():
        baseline_key = str(key)
        if not baseline_key.lower().startswith("baseline_") or not valid(baseline):
            continue
        baseline_name = baseline_key[len("Baseline_") :]
        target = _normalized_metric_name(baseline_name)
        candidates = []
        for candidate_key, candidate_value in metrics.items():
            candidate_name = str(candidate_key)
            if candidate_name == baseline_key or not valid(candidate_value):
                continue
            direct = _normalized_metric_name(candidate_name)
            stripped = _normalized_metric_name(_strip_metric_group_prefix(candidate_name))
            if direct == target or stripped == target:
                candidates.append((0 if direct == target else 1, candidate_name))
        if not candidates:
            continue
        experiment_key = min(candidates, key=lambda item: (item[0], item[1]))[1]
        display_name = _strip_metric_group_prefix(experiment_key).replace("_", " ").strip()
        pairs.append((display_name or baseline_name, baseline_key, experiment_key))
    return pairs


def _preferred_metric_pairs(
    pairs: list[tuple[str, str, str]],
) -> list[tuple[str, str, str]]:
    def priority(pair: tuple[str, str, str]) -> tuple[int, str]:
        name = _normalized_metric_name(pair[0])
        if "accuracy" in name:
            return (0, name)
        if any(token in name for token in ("error", "f1", "auc", "precision", "recall")):
            return (1, name)
        if "loss" in name:
            return (3, name)
        return (2, name)

    return sorted(pairs, key=priority)


def _deduplicated_metrics(metrics: dict[str, Any]) -> list[tuple[str, Any]]:
    ordered = sorted(
        metrics.items(),
        key=lambda item: (
            1 if _strip_metric_group_prefix(str(item[0])) != str(item[0]) else 0,
            str(item[0]),
        ),
    )
    seen = set()
    kept = []
    for key, value in ordered:
        name = str(key)
        role = "baseline" if name.lower().startswith("baseline_") else "experiment"
        canonical = _normalized_metric_name(
            name[len("Baseline_") :] if role == "baseline" else _strip_metric_group_prefix(name)
        )
        marker = (role, canonical)
        if marker in seen:
            continue
        seen.add(marker)
        kept.append((name, value))
    return kept


def _strip_metric_group_prefix(value: str) -> str:
    return re.sub(
        r"^(?:ls|label[_\s-]*smoothing|variant|experiment|treatment)[_\s-]+",
        "",
        str(value),
        flags=re.IGNORECASE,
    )


def _normalized_metric_name(value: str) -> str:
    return re.sub(r"[^a-z0-9\u4e00-\u9fff]+", "", str(value).lower())


def _pick(value: dict, *keys: str) -> Any:
    for key in keys:
        if key in value and value[key] not in (None, "", [], {}):
            return value[key]
    return None


def _short_text(value: Any, limit: int) -> str:
    if isinstance(value, dict):
        text = "；".join(
            f"{_human_label(key)}：{_short_text(item, 180)}"
            for key, item in _curated_items(value)
            if item not in (None, "", [], {})
        )
    elif isinstance(value, list):
        text = "；".join(_short_text(item, 180) for item in value)
    else:
        text = str(value)
    return text if len(text) <= limit else text[: limit - 1].rstrip() + "…"


def _human_label(value: object) -> str:
    raw = str(value)
    labels = {
        "name": "名称",
        "source": "来源",
        "source_type": "来源类型",
        "split": "数据划分",
        "preprocessing": "预处理",
        "limitations": "已知限制",
        "availability": "可用状态",
        "license": "许可",
        "description": "说明",
        "learning_rate": "学习率",
        "batch_size": "批量大小",
        "epochs": "训练轮数",
        "optimizer": "优化器",
        "label_smoothing_epsilon": "标签平滑系数 ε",
        "patience": "早停耐心值",
        "experiment_id": "实验编号",
        "provider": "运行方式",
        "parameters": "关键参数",
        "seeds": "随机种子",
        "environment": "运行环境",
        "summary": "结果概述",
        "conclusion": "实验结论",
        "interpretation": "结果分析",
        "Overall Test Accuracy": "整体测试准确率",
        "Test Accuracy": "测试准确率",
        "Final Training Loss": "最终训练损失",
        "Target Class Confusion Error Rate": "目标类别混淆错误率",
        "Total Parameters": "参数量",
    }
    return labels.get(raw, raw.replace("_", " ").strip())


def _format_number(value: Any) -> str:
    if isinstance(value, float):
        return f"{value:.6g}"
    return str(value)


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_run_font(
    run,
    name: str,
    size: float,
    color: RGBColor | None = None,
    bold: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
