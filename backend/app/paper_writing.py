from __future__ import annotations

from collections.abc import Callable
from datetime import datetime, timezone
from io import BytesIO
import json
import re
from threading import Lock, Thread
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from backend.app.reporting import attachment_headers
from backend.app.providers.llm import LLMRequestCancelled
from backend.app.storage.repository import Repository
from backend.app.workflow.skills import SkillLoader


RUNNING_STATES = {"queued", "planning", "writing", "auditing", "revising"}


class PaperWritingStopped(RuntimeError):
    pass


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


class PaperWritingManager:
    def __init__(
        self,
        repository: Repository,
        llm_provider: Callable[[], object],
        skill_loader: SkillLoader,
    ) -> None:
        self.repository = repository
        self._llm_provider = llm_provider
        self.skill_loader = skill_loader
        self._threads: dict[str, Thread] = {}
        self._guard = Lock()

    def get(self, run_id: str) -> dict:
        run = self.repository.get_run(run_id)
        return run.paper_writing or self._empty_state()

    def start(self, run_id: str, config: dict) -> dict:
        run = self.repository.get_run(run_id)
        if not any(item.type == "report" for item in run.artifacts):
            raise ValueError("PAPER_REPORT_REQUIRED")
        current = run.paper_writing or {}
        if current.get("status") in RUNNING_STATES:
            raise ValueError("PAPER_WRITING_ALREADY_RUNNING")
        state = {
            **self._empty_state(),
            "status": "queued",
            "stage": "准备论文规划",
            "progress": 2,
            "config": {
                "venue": str(config.get("venue") or "未指定"),
                "language": str(config.get("language") or "zh-CN"),
                "paper_type": str(config.get("paper_type") or "实验研究论文"),
                "authors": str(config.get("authors") or ""),
                "notes": str(config.get("notes") or ""),
            },
            "started_at": utc_now(),
            "updated_at": utc_now(),
            "stop_requested": False,
            "error": "",
            "plan": {},
            "sections": [],
            "audit": {},
        }
        self.repository.update_paper_writing(run_id, state)
        self._launch(run_id, self._generate_plan)
        return self.get(run_id)

    def confirm_plan(self, run_id: str, feedback: str = "") -> dict:
        state = self.get(run_id)
        if state.get("status") not in {"waiting_plan_confirmation", "interrupted", "failed"}:
            raise ValueError("PAPER_PLAN_NOT_WAITING")
        self.repository.update_paper_writing(
            run_id,
            {
                "status": "queued",
                "stage": "根据确认的大纲准备写作",
                "progress": max(20, int(state.get("progress") or 0)),
                "plan_feedback": feedback.strip(),
                "stop_requested": False,
                "error": "",
                "updated_at": utc_now(),
            },
        )
        if not (state.get("plan") or {}).get("sections"):
            target = self._generate_plan
        else:
            target = self._regenerate_plan if feedback.strip() else self._write_sections
        self._launch(run_id, target)
        return self.get(run_id)

    def finalize(self, run_id: str, feedback: str = "") -> dict:
        state = self.get(run_id)
        if state.get("status") != "waiting_final_confirmation":
            raise ValueError("PAPER_FINAL_NOT_WAITING")
        if feedback.strip():
            self.repository.update_paper_writing(
                run_id,
                {
                    "status": "queued",
                    "stage": "根据反馈修订论文",
                    "progress": 82,
                    "revision_feedback": feedback.strip(),
                    "stop_requested": False,
                    "error": "",
                    "updated_at": utc_now(),
                },
            )
            self._launch(run_id, self._revise_sections)
        else:
            self.repository.update_paper_writing(
                run_id,
                {
                    "status": "completed",
                    "stage": "论文写作完成",
                    "progress": 100,
                    "completed_at": utc_now(),
                    "updated_at": utc_now(),
                    "active_skill": "",
                },
            )
        return self.get(run_id)

    def stop(self, run_id: str) -> dict:
        state = self.get(run_id)
        if state.get("status") not in RUNNING_STATES:
            return state
        self.repository.update_paper_writing(
            run_id,
            {
                "stop_requested": True,
                "stage": "正在停止论文写作",
                "updated_at": utc_now(),
            },
        )
        cancel = getattr(self._llm_provider(), "cancel_run", None)
        if callable(cancel):
            cancel(run_id)
        return self.get(run_id)

    def recover(self) -> list[str]:
        interrupted = []
        for run in self.repository.list_runs():
            if (run.paper_writing or {}).get("status") in RUNNING_STATES:
                self.repository.update_paper_writing(
                    run.id,
                    {
                        "status": "interrupted",
                        "stage": "后端重启，论文写作已暂停",
                        "error": "可以从当前确认节点继续。",
                        "stop_requested": False,
                        "updated_at": utc_now(),
                    },
                )
                interrupted.append(run.id)
        return interrupted

    def word_bytes(self, run_id: str) -> bytes:
        run = self.repository.get_run(run_id)
        state = run.paper_writing or {}
        if state.get("status") != "completed":
            raise ValueError("PAPER_NOT_COMPLETED")
        return build_paper_docx(state)

    def latex_bytes(self, run_id: str) -> bytes:
        run = self.repository.get_run(run_id)
        state = run.paper_writing or {}
        if state.get("status") != "completed":
            raise ValueError("PAPER_NOT_COMPLETED")
        report = next(item.content for item in reversed(run.artifacts) if item.type == "report")
        return build_latex_package(state, report)

    def _launch(self, run_id: str, target) -> None:
        with self._guard:
            current = self._threads.get(run_id)
            if current and current.is_alive():
                raise ValueError("PAPER_WRITING_ALREADY_RUNNING")
            thread = Thread(
                target=self._run_guarded,
                args=(run_id, target),
                name=f"paper-writing-{run_id}",
                daemon=True,
            )
            self._threads[run_id] = thread
            thread.start()

    def _run_guarded(self, run_id: str, target) -> None:
        provider = self._llm_provider()
        begin = getattr(provider, "begin_run", None)
        end = getattr(provider, "end_run", None)
        try:
            if callable(begin):
                begin(run_id)
            target(run_id)
        except (PaperWritingStopped, LLMRequestCancelled):
            self.repository.update_paper_writing(
                run_id,
                {
                    "status": "interrupted",
                    "stage": "论文写作已暂停",
                    "error": "",
                    "active_skill": "",
                    "stop_requested": False,
                    "updated_at": utc_now(),
                },
            )
        except Exception as exc:
            self.repository.update_paper_writing(
                run_id,
                {
                    "status": "failed",
                    "stage": "论文写作失败",
                    "error": str(exc),
                    "active_skill": "",
                    "updated_at": utc_now(),
                },
            )
        finally:
            if callable(end):
                end(run_id)
            with self._guard:
                self._threads.pop(run_id, None)

    def _generate_plan(self, run_id: str) -> None:
        self._ensure_not_stopped(run_id)
        self._update(
            run_id,
            status="planning",
            stage="Qwen 正在分析研究报告并规划论文",
            progress=8,
            active_skill="paper-writing + paper-plan",
        )
        run, report = self._run_report(run_id)
        instructions = self._skills("paper-writing", "paper-plan")
        config = self.get(run_id).get("config") or {}
        plan = self._llm_provider().generate_json(
            "paper.plan",
            {
                "report": _compact_report(report),
                "experiment_results": [
                    item.content for item in run.artifacts if item.type == "experiment_result"
                ],
                "config": config,
            },
            {
                "title": "string",
                "research_question": "string",
                "terminal_verdict": "string",
                "contributions": ["string"],
                "sections": [
                    {
                        "id": "string",
                        "title": "string",
                        "purpose": "string",
                        "key_points": ["string"],
                        "evidence": ["string"],
                        "citations": ["string"],
                    }
                ],
                "claims_evidence": [{"claim": "string", "evidence": ["string"]}],
                "figures": [{"title": "string", "data_source": "string", "required": "boolean"}],
                "limitations": ["string"],
            },
            instructions=(
                f"{instructions}\n\n"
                "当前执行者是 Qwen。严格输出符合 schema 的 JSON。"
                f"写作语言为 {config.get('language', 'zh-CN')}，投稿方向为 {config.get('venue', '未指定')}。"
            ),
        )
        self._ensure_not_stopped(run_id)
        self._update(
            run_id,
            status="waiting_plan_confirmation",
            stage="论文大纲等待确认",
            progress=20,
            active_skill="paper-plan",
            plan=plan,
            references=report.get("References") or report.get("参考文献") or [],
            completed_sections=0,
            total_sections=len(plan.get("sections") or []),
        )

    def _regenerate_plan(self, run_id: str) -> None:
        state = self.get(run_id)
        self._update(
            run_id,
            status="planning",
            stage="Qwen 正在根据反馈修改论文大纲",
            progress=12,
            active_skill="paper-plan",
        )
        run, report = self._run_report(run_id)
        plan = self._llm_provider().generate_json(
            "paper.plan",
            {
                "report": _compact_report(report),
                "experiment_results": [
                    item.content for item in run.artifacts if item.type == "experiment_result"
                ],
                "current_plan": state.get("plan") or {},
                "human_feedback": state.get("plan_feedback") or "",
                "config": state.get("config") or {},
            },
            {
                "title": "string",
                "research_question": "string",
                "terminal_verdict": "string",
                "contributions": ["string"],
                "sections": ["object"],
                "claims_evidence": ["object"],
                "figures": ["object"],
                "limitations": ["string"],
            },
            instructions=self._skills("paper-writing", "paper-plan"),
        )
        self._update(
            run_id,
            status="waiting_plan_confirmation",
            stage="修改后的论文大纲等待确认",
            progress=20,
            active_skill="paper-plan",
            plan=plan,
            total_sections=len(plan.get("sections") or []),
        )

    def _write_sections(self, run_id: str) -> None:
        state = self.get(run_id)
        plan = state.get("plan") or {}
        planned_sections = plan.get("sections") or []
        if not planned_sections:
            raise ValueError("PAPER_PLAN_HAS_NO_SECTIONS")
        run, report = self._run_report(run_id)
        instructions = self._skills("paper-writing", "paper-write")
        references = report.get("References") or report.get("参考文献") or []
        experiment_results = [
            item.content for item in run.artifacts if item.type == "experiment_result"
        ]
        completed = []
        for index, section in enumerate(planned_sections):
            self._ensure_not_stopped(run_id)
            title = str(section.get("title") or f"Section {index + 1}")
            progress = 24 + int(52 * index / max(len(planned_sections), 1))
            self._update(
                run_id,
                status="writing",
                stage=f"Qwen 正在撰写：{title}",
                current_section=title,
                progress=progress,
                active_skill="paper-write",
                completed_sections=len(completed),
                total_sections=len(planned_sections),
            )
            drafted = self._llm_provider().generate_json(
                "paper.write_section",
                {
                    "approved_plan": plan,
                    "section": section,
                    "completed_sections": completed,
                    "report": _compact_report(report),
                    "experiment_results": experiment_results,
                    "verified_references": references,
                    "config": state.get("config") or {},
                },
                {"id": "string", "title": "string", "content": "string", "citations": ["string"]},
                instructions=(
                    f"{instructions}\n\n只写当前章节。"
                    "不得输出内部路径、artifact ID、哈希或不存在的引用。"
                ),
            )
            completed.append(drafted)
            self._update(
                run_id,
                sections=completed,
                completed_sections=len(completed),
                progress=24 + int(52 * len(completed) / len(planned_sections)),
            )
        self._audit(run_id)

    def _revise_sections(self, run_id: str) -> None:
        state = self.get(run_id)
        sections = state.get("sections") or []
        feedback = state.get("revision_feedback") or ""
        run, report = self._run_report(run_id)
        revised = []
        instructions = self._skills("paper-writing", "paper-write")
        for index, section in enumerate(sections):
            self._ensure_not_stopped(run_id)
            title = str(section.get("title") or f"Section {index + 1}")
            self._update(
                run_id,
                status="revising",
                stage=f"Qwen 正在修订：{title}",
                current_section=title,
                active_skill="paper-write",
                progress=82 + int(10 * index / max(len(sections), 1)),
            )
            value = self._llm_provider().generate_json(
                "paper.revise_section",
                {
                    "section": section,
                    "human_feedback": feedback,
                    "audit": state.get("audit") or {},
                    "report": _compact_report(report),
                    "experiment_results": [
                        item.content for item in run.artifacts if item.type == "experiment_result"
                    ],
                },
                {"id": "string", "title": "string", "content": "string", "citations": ["string"]},
                instructions=instructions,
            )
            revised.append(value)
            self._update(run_id, sections=revised + sections[index + 1 :])
        self._audit(run_id)

    def _audit(self, run_id: str) -> None:
        self._ensure_not_stopped(run_id)
        state = self.get(run_id)
        run, report = self._run_report(run_id)
        self._update(
            run_id,
            status="auditing",
            stage="Qwen 正在核对实验数值、引用与结论边界",
            progress=94,
            active_skill="paper-writing",
            current_section="",
        )
        audit = self._llm_provider().generate_json(
            "paper.audit",
            {
                "plan": state.get("plan") or {},
                "sections": state.get("sections") or [],
                "experiment_results": [
                    item.content for item in run.artifacts if item.type == "experiment_result"
                ],
                "verified_references": report.get("References") or report.get("参考文献") or [],
                "terminal_report": _compact_report(report),
            },
            {
                "accepted": "boolean",
                "summary": "string",
                "issues": ["string"],
                "numeric_claim_checks": ["object"],
                "citation_checks": ["object"],
            },
            instructions=self._skills("paper-writing"),
        )
        self._update(
            run_id,
            status="waiting_final_confirmation",
            stage="论文初稿与审计结果等待确认",
            progress=98,
            active_skill="paper-writing",
            audit=audit,
        )

    def _ensure_not_stopped(self, run_id: str) -> None:
        if self.get(run_id).get("stop_requested") is True:
            self._update(
                run_id,
                status="interrupted",
                stage="论文写作已暂停",
                active_skill="",
                stop_requested=False,
            )
            raise PaperWritingStopped("PAPER_WRITING_STOPPED")

    def _run_report(self, run_id: str):
        run = self.repository.get_run(run_id)
        report = next(
            (item.content for item in reversed(run.artifacts) if item.type == "report"),
            None,
        )
        if report is None:
            raise ValueError("PAPER_REPORT_REQUIRED")
        return run, report

    def _skills(self, *skill_ids: str) -> str:
        return "\n\n".join(
            f"## Skill: {skill_id}\n{self.skill_loader.load_complete(skill_id).instructions}"
            for skill_id in skill_ids
        )

    def _update(self, run_id: str, **values) -> None:
        values["updated_at"] = utc_now()
        self.repository.update_paper_writing(run_id, values)

    @staticmethod
    def _empty_state() -> dict:
        return {
            "status": "not_started",
            "stage": "尚未开始",
            "progress": 0,
            "config": {},
            "plan": {},
            "sections": [],
            "audit": {},
            "references": [],
            "active_skill": "",
            "current_section": "",
            "completed_sections": 0,
            "total_sections": 0,
            "stop_requested": False,
            "error": "",
            "started_at": "",
            "updated_at": "",
            "completed_at": "",
        }


def build_paper_docx(state: dict) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)
    normal = document.styles["Normal"]
    normal.font.name = "宋体" if state.get("config", {}).get("language") == "zh-CN" else "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(7)
    for style_name, size in (("Title", 20), ("Heading 1", 14), ("Heading 2", 12)):
        style = document.styles[style_name]
        style.font.name = "微软雅黑" if state.get("config", {}).get("language") == "zh-CN" else "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 78, 121)

    plan = state.get("plan") or {}
    title = str(plan.get("title") or "研究论文")
    paragraph = document.add_paragraph(style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(title)
    authors = str((state.get("config") or {}).get("authors") or "")
    if authors:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(authors)
    venue = str((state.get("config") or {}).get("venue") or "")
    if venue:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(venue)
        run.font.color.rgb = RGBColor(100, 116, 139)
    for section_value in state.get("sections") or []:
        document.add_heading(str(section_value.get("title") or ""), level=1)
        for block in re.split(r"\n\s*\n", str(section_value.get("content") or "")):
            if block.strip():
                p = document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Inches(0.28)
                p.add_run(block.strip())
    references = state.get("references") or []
    if references:
        heading = "参考文献" if state.get("config", {}).get("language") == "zh-CN" else "References"
        document.add_heading(heading, level=1)
        for index, reference in enumerate(references, 1):
            if not isinstance(reference, dict):
                continue
            authors = reference.get("authors") or []
            if isinstance(authors, list):
                authors = ", ".join(str(item) for item in authors)
            identifiers = reference.get("identifiers") if isinstance(reference.get("identifiers"), dict) else {}
            identity = identifiers.get("doi") or identifiers.get("arxiv") or reference.get("url") or ""
            text = " ".join(
                str(item)
                for item in (
                    authors,
                    f"({reference.get('year')})." if reference.get("year") else "",
                    reference.get("title"),
                    identity,
                )
                if item
            )
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.28)
            p.paragraph_format.first_line_indent = Inches(-0.28)
            p.add_run(f"[{index}] {text}")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_latex_package(state: dict, report: dict) -> bytes:
    sections = state.get("sections") or []
    config = state.get("config") or {}
    chinese = config.get("language") == "zh-CN"
    title = _latex_escape(str((state.get("plan") or {}).get("title") or "Research Paper"))
    authors = _latex_escape(str(config.get("authors") or ""))
    document_class = "\\documentclass[UTF8,11pt]{ctexart}" if chinese else "\\documentclass[11pt]{article}"
    includes = []
    stream = BytesIO()
    with ZipFile(stream, "w", ZIP_DEFLATED) as archive:
        for index, section in enumerate(sections, 1):
            filename = f"section_{index:02d}.tex"
            includes.append(f"\\input{{sections/{filename}}}")
            content = (
                f"\\section{{{_latex_escape(str(section.get('title') or ''))}}}\n\n"
                f"{_latex_prose(str(section.get('content') or ''))}\n"
            )
            archive.writestr(f"sections/{filename}", content.encode("utf-8"))
        bibliography = _bibtex(report.get("References") or report.get("参考文献") or [])
        main = (
            f"{document_class}\n"
            "\\usepackage{geometry}\n\\usepackage{booktabs}\n\\usepackage{hyperref}\n"
            "\\geometry{margin=1in}\n"
            f"\\title{{{title}}}\n\\author{{{authors}}}\n\\date{{}}\n"
            "\\begin{document}\n\\maketitle\n\n"
            + "\n".join(includes)
            + "\n\\bibliographystyle{plain}\n\\bibliography{references}\n\\end{document}\n"
        )
        archive.writestr("main.tex", main.encode("utf-8"))
        archive.writestr("references.bib", bibliography.encode("utf-8"))
    return stream.getvalue()


def _compact_report(report: dict) -> dict:
    keys = (
        "Paper Title", "Paper Abstract", "Executive Summary", "Problem Statement",
        "Rationale", "Technical Details", "Datasets", "Source", "Target", "Methods",
        "Experiments", "Iteration Summary", "Results", "Report Evidence", "Research Conclusion",
        "Limitations", "References", "标题", "摘要", "执行摘要", "待研究问题",
        "解决思路", "方法论", "实验设计", "实验结果", "研究结论", "参考文献",
    )
    return {key: report[key] for key in keys if key in report}


def _latex_escape(value: str) -> str:
    replacements = {
        "\\": "\\textbackslash{}",
        "&": "\\&",
        "%": "\\%",
        "$": "\\$",
        "#": "\\#",
        "_": "\\_",
        "{": "\\{",
        "}": "\\}",
        "~": "\\textasciitilde{}",
        "^": "\\textasciicircum{}",
    }
    return "".join(replacements.get(char, char) for char in value)


def _latex_prose(value: str) -> str:
    return "\n\n".join(_latex_escape(block.strip()) for block in re.split(r"\n\s*\n", value) if block.strip())


def _bibtex(references: list) -> str:
    entries = []
    used = set()
    for index, reference in enumerate(references, 1):
        if not isinstance(reference, dict):
            continue
        authors = reference.get("authors") or []
        if isinstance(authors, list):
            authors = " and ".join(str(item) for item in authors)
        year = str(reference.get("year") or "")
        title = str(reference.get("title") or f"Reference {index}")
        base = re.sub(r"[^A-Za-z0-9]", "", str(authors).split(" ")[-1] + year) or f"ref{index}"
        key = base
        suffix = 1
        while key in used:
            suffix += 1
            key = f"{base}{suffix}"
        used.add(key)
        identifiers = reference.get("identifiers") if isinstance(reference.get("identifiers"), dict) else {}
        fields = [
            f"  title = {{{title}}}",
            f"  author = {{{authors}}}",
            f"  year = {{{year}}}",
        ]
        if identifiers.get("doi"):
            fields.append(f"  doi = {{{identifiers['doi']}}}")
        if reference.get("url"):
            fields.append(f"  url = {{{reference['url']}}}")
        entries.append(f"@article{{{key},\n" + ",\n".join(fields) + "\n}")
    return "\n\n".join(entries) + "\n"
