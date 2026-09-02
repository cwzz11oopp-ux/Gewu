from docx import Document
from io import BytesIO
import pytest

from backend.app.models.artifact import Artifact
from backend.app.report_visualization import build_report_spec, render_figure_png
from backend.app.reporting import FIGURE_SECTION, build_report_docx
from backend.app.workflow.research_synthesis import stable_paper_id


def artifact(kind: str, content: dict, identifier: str) -> Artifact:
    return Artifact(
        id=identifier,
        run_id="run_visual",
        type=kind,
        version=1,
        title=kind,
        content=content,
        source_step=kind,
        created_by="test",
    )


def test_seed_delta_figure_uses_the_same_error_direction_as_statistics():
    from backend.app.report_visualization import _seed_delta_figure
    for metric in ("mae", "brier_score", "ece"):
        figure = _seed_delta_figure({"seed_results": [
            {"seed": 1, "metrics": {f"baseline_{metric}": .3, f"variant_{metric}": .2}},
            {"seed": 2, "metrics": {f"baseline_{metric}": .4, f"variant_{metric}": .2}},
        ]}, None)
        assert figure.chart.metric_direction == "lower"


def test_saved_legacy_arxiv_citations_resolve_without_allowing_unknown_sources():
    from backend.app.reporting import _render_citations
    report = {"References": [{"title": "Existing source", "identifiers": {"arxiv": "2101.12445"}}]}
    assert _render_citations("旧引用[PAPER-2101.12445v1]。", report) == "旧引用[1]。"
    with pytest.raises(ValueError, match="REPORT_CITATION_UNRESOLVED"):
        _render_citations("[PAPER-9999.12345v1]", report)


def test_docx_renders_all_cited_references_and_saved_statistics_without_rewriting():
    refs = [{"title": f"Source {index}", "identifiers": {"doi": f"10.1/{index}"}} for index in range(16)]
    identifier = stable_paper_id(refs[-1])
    report = {
        "Paper Abstract": f"摘要引用[{identifier}]。",
        "Narrative Sections": [{"id": "results", "title": "结果", "paragraphs": [f"论点[{identifier}]。"]}],
        "References": refs,
        "Report Evidence": {"deterministic_result_evidence": {
            "metric": "AUC", "mean_delta": -.02689279047,
            "paired_t_test": {"method": "paired_t_test", "status": "computed", "n_pairs": 3,
                              "statistic": -3.767492203, "p_value": .06378557524},
            "confidence_interval_95": [-.0576056, .0038200], "confidence_interval_method": "student_t",
        }},
    }
    doc = Document(BytesIO(build_report_docx(report, run_id="test", run_title="test")))
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "摘要引用[16]" in text and "论点[16]" in text
    assert "[16] Source 15" in text
    assert "PAPER-" not in text
    tables = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    assert "0.06378557524" in tables and "-3.767492203" in tables
    assert "0.227" not in tables


def test_report_spec_uses_persisted_values_and_never_invents_training_curve():
    plan = artifact("plan", {"parameters": {"learning_rate": 0.001, "batch_size": 64}, "seeds": [1, 2]}, "art_plan")
    result = artifact(
        "experiment_result",
        {
            "is_real_experiment": True,
            "metrics": {"CNN Test Accuracy": 91.362, "MLP Test Accuracy": 88.888},
            "seed_results": [
                {"seed": 1, "CNN Test Accuracy": 91.0, "MLP Test Accuracy": 88.1},
                {"seed": 2, "CNN Test Accuracy": 91.7, "MLP Test Accuracy": 89.7},
            ],
        },
        "art_result",
    )
    spec = build_report_spec(
        [plan, result],
        selected_figure_ids=["research_workflow", "control_variables", "seed_comparison", "main_comparison", "training_curve", "workflow_timeline"],
        decision_rationale="仅选择已保存数据支持的图。",
    )

    figures = {item["figure_id"]: item for item in spec["figures"]}
    assert figures["main_comparison"]["source_artifact_ids"] == ["art_result"]
    assert figures["main_comparison"]["chart"]["series"][0]["values"] == [91.362, 88.888]
    assert figures["seed_comparison"]["chart"]["labels"] == ["1", "2"]
    assert "training_curve" not in figures
    assert any(item["figure_id"] == "training_curve" and "epoch/step" in item["reason"] for item in spec["omitted_figures"])
    assert render_figure_png(figures["main_comparison"]).startswith(b"\x89PNG")


def test_visual_anchors_are_selected_and_distributed_before_results():
    plan = artifact(
        "plan",
        {
            "comparisons": [{"baseline": "CNN", "variant": "CNN + Label Smoothing"}],
            "parameters": {"learning_rate": 0.001, "batch_size": 64},
            "seeds": [1, 2],
            "procedure": {"steps": ["train", "evaluate"]},
        },
        "art_plan",
    )
    result = artifact(
        "experiment_result",
        {
            "metrics": {"Baseline_Test_Accuracy": 0.88, "LS_Test_Accuracy": 0.89},
            "seed_results": [
                {"seed": 1, "Baseline_Test_Accuracy": 0.87, "LS_Test_Accuracy": 0.89},
                {"seed": 2, "Baseline_Test_Accuracy": 0.89, "LS_Test_Accuracy": 0.90},
            ],
            "epoch_metrics": [
                {"epoch": 1, "loss": 1.2},
                {"epoch": 2, "loss": 0.8},
            ],
        },
        "art_result",
    )

    spec = build_report_spec(
        [plan, result],
        selected_figure_ids=[
            "model_structure", "method_pipeline", "control_variables", "workflow_timeline",
            "training_curve", "main_comparison", "seed_comparison", "seed_delta",
        ],
    )

    figure_ids = [item["figure_id"] for item in spec["figures"]]
    assert figure_ids[:5] == [
        "model_structure", "method_pipeline", "control_variables", "workflow_timeline", "training_curve",
    ]
    assert [FIGURE_SECTION[figure_id] for figure_id in figure_ids] == [
        "method", "design", "design", "iteration", "iteration", "results", "results", "results",
    ]


def test_docx_embeds_only_report_spec_figures_and_lists_missing_data_reason():
    result = artifact("experiment_result", {"metrics": {"A": 0.8, "B": 0.9}}, "art_result")
    spec = build_report_spec([result], selected_figure_ids=["main_comparison"], decision_rationale="主比较图即可。")
    payload = build_report_docx(
        {"Paper Title": "可视化报告", "Paper Abstract": "摘要", "Report Spec": spec},
        run_id="run_visual",
        run_title="可视化报告",
    )
    document = Document(__import__("io").BytesIO(payload))
    assert len(document.inline_shapes) == 1
    assert any("未生成 training_curve" in paragraph.text for paragraph in document.paragraphs)


def test_model_comparison_figure_uses_frozen_plan_and_never_invents_eca():
    plan = artifact(
        "plan",
        {
            "method": {
                "name": "固定预算下的标签平滑实验",
                "components": ["仅修改损失函数"],
            },
            "comparisons": [
                {
                    "baseline": "Standard Cross-Entropy Loss",
                    "variant": "Label Smoothing Loss (epsilon=0.1)",
                }
            ],
        },
        "art_plan",
    )

    spec = build_report_spec(
        [plan],
        selected_figure_ids=["model_structure"],
    )

    figure = next(item for item in spec["figures"] if item["figure_id"] == "model_structure")
    serialized = str(figure)
    assert "Standard Cross-Entropy Loss" in serialized
    assert "Label Smoothing Loss (epsilon=0.1)" in serialized
    assert "ECA" not in serialized
    assert "四层卷积" not in serialized
    assert "浅层特征" not in serialized


def test_model_comparison_figure_is_omitted_without_explicit_comparison():
    plan = artifact(
        "plan",
        {"method": {"name": "ECA method", "components": ["ECA"]}},
        "art_plan",
    )

    spec = build_report_spec([plan], selected_figure_ids=["model_structure"])

    assert not any(item["figure_id"] == "model_structure" for item in spec["figures"])


def test_harness_metric_aliases_produce_accuracy_and_seed_delta_figures():
    result = artifact(
        "experiment_result",
        {
            "metrics": {
                "Baseline_Test_Accuracy": 0.883,
                "LS_Test_Accuracy": 0.8857,
                "Test Accuracy": 0.8857,
                "Baseline_Final_Training_Loss": 0.2736,
                "Final Training Loss": 0.7408,
            },
            "seed_results": [
                {
                    "seed": 1,
                    "metrics": {
                        "Baseline_Test_Accuracy": 0.8834,
                        "LS_Test_Accuracy": 0.8889,
                        "Test Accuracy": 0.8889,
                    },
                },
                {
                    "seed": 2,
                    "metrics": {
                        "Baseline_Test_Accuracy": 0.8822,
                        "LS_Test_Accuracy": 0.8849,
                        "Test Accuracy": 0.8849,
                    },
                },
            ],
        },
        "art_result",
    )

    spec = build_report_spec(
        [result],
        selected_figure_ids=["main_comparison", "seed_comparison", "seed_delta"],
    )
    figures = {item["figure_id"]: item for item in spec["figures"]}

    assert figures["main_comparison"]["chart"]["labels"] == ["测试准确率"]
    assert figures["main_comparison"]["chart"]["series"][0]["values"] == [0.883]
    assert figures["main_comparison"]["chart"]["series"][1]["values"] == [0.8857]
    assert figures["seed_comparison"]["chart"]["series"][1]["values"] == [0.8889, 0.8849]
    assert figures["seed_delta"]["chart"]["series"][0]["values"] == [
        0.8889 - 0.8834,
        0.8849 - 0.8822,
    ]


def test_narrative_support_tables_hide_planning_metadata_and_deduplicate_aliases():
    report = {
        "Paper Title": "Label Smoothing 报告",
        "Paper Abstract": "摘要",
        "Narrative Sections": [
            {"id": "design", "title": "实验设计", "paragraphs": ["设计正文"]},
            {"id": "results", "title": "实验结果", "paragraphs": ["结果正文"]},
            {"id": "conclusion", "title": "研究结论", "paragraphs": ["结论正文"]},
        ],
        "Experiments": {
            "parameters": {
                "epochs": 5,
                "optimizer": "SGD",
                "label_smoothing_epsilon": 0.1,
                "additional_sections": {
                    "training_budget_rationale": "unverified planning rationale"
                },
            },
            "seeds": [1, 2],
        },
        "Results": {
            "metric_summary": {
                "Baseline_Test_Accuracy": {"mean": 0.883, "std": 0.001},
                "LS_Test_Accuracy": {"mean": 0.886, "std": 0.002},
                "Test Accuracy": {"mean": 0.886, "std": 0.002},
            },
            "seed_results": [
                {
                    "seed": 1,
                    "metrics": {
                        "Baseline_Test_Accuracy": 0.883,
                        "LS_Test_Accuracy": 0.887,
                        "Test Accuracy": 0.887,
                    },
                },
                {
                    "seed": 2,
                    "metrics": {
                        "Baseline_Test_Accuracy": 0.882,
                        "LS_Test_Accuracy": 0.885,
                        "Test Accuracy": 0.885,
                    },
                },
            ],
        },
        "Reproducibility": {
            "experiment_id": "experiment_1",
            "parameters": {
                "epochs": 5,
                "optimizer": "SGD",
                "label_smoothing_epsilon": 0.1,
                "additional_sections": {
                    "training_budget_rationale": "unverified planning rationale"
                },
            },
            "seeds": [1, 2],
        },
    }

    document = Document(
        __import__("io").BytesIO(
            build_report_docx(report, run_id="run_visual", run_title="Label Smoothing 报告")
        )
    )
    text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )

    assert "additional sections" not in text.lower()
    assert "training budget rationale" not in text.lower()
    assert "unverified planning rationale" not in text
    assert "标签平滑系数 ε" in text
    assert "LS Test Accuracy" not in text
    assert "逐随机种子的测试准确率配对结果" in text
